#!/usr/bin/env python3
from __future__ import annotations

import os, csv, math
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------
# Paths
# ----------------------------
I9 = Path(os.environ.get("I9", "/mnt/e/Bacillus WGS/bacillus_paranthracis/06_reports_wsl/I-9"))
OUTDIR = Path(os.environ.get("OUTDIR", str(I9 / "figures" / "_paper_mainfigs" / "tables")))
OUTDIR.mkdir(parents=True, exist_ok=True)

K_READS   = I9 / "I-9.reads.kraken2.report.txt"
K_CONTIG  = I9 / "I-9.contigs.kraken2.report.txt"
K_MIN5    = I9 / "I-9.contigs.min5kb.kraken2.report.txt"

FA_ALL    = I9 / "I-9.final.contigs.fa"
FA_MIN5   = I9 / "I-9.contigs.min5kb.fa"
FA_BAC    = I9 / "qc_clean" / "I-9.bacillus_only.fa"

TBL_DIR   = I9 / "tables"
TBL_BAC   = I9 / "tables_bacillus_only"

# ----------------------------
# Helpers
# ----------------------------
def read_tsv(path: Path) -> list[dict]:
    with path.open("r", newline="") as f:
        rdr = csv.DictReader(f, delimiter="\t")
        return [dict(r) for r in rdr]

def try_paths(paths: list[Path]) -> Path | None:
    for p in paths:
        if p and p.exists() and p.is_file() and p.stat().st_size > 0:
            return p
    return None

def parse_kraken_report(path: Path) -> list[dict]:
    rows = []
    if not path.exists():
        return rows
    with path.open("r", encoding="utf-8", errors="replace") as f:
        for line in f:
            if not line.strip():
                continue
            parts = line.rstrip("\n").split("\t")
            if len(parts) < 6:
                continue
            rows.append({
                "pct": float(parts[0].strip()),
                "rank": parts[3].strip(),
                "name": parts[5].strip(),
            })
    return rows

def kraken_genus_distribution(rows: list[dict]) -> dict[str, float]:
    dist = {}
    for r in rows:
        if r["rank"] == "G":
            dist[r["name"]] = dist.get(r["name"], 0.0) + r["pct"]
        if r["rank"] == "U" and r["name"].lower().startswith("unclassified"):
            dist["Unclassified"] = r["pct"]
    return dist

def diversity_from_pct(dist: dict[str, float]) -> tuple[float, float]:
    ps = []
    for v in dist.values():
        p = max(v, 0.0) / 100.0
        if p > 0:
            ps.append(p)
    if not ps:
        return (0.0, 0.0)
    H = -sum(p * math.log(p) for p in ps)
    D = sum(p*p for p in ps)
    return (H, 1.0 - D)

def fasta_stats(path: Path) -> dict:
    if not path.exists() or path.stat().st_size == 0:
        return {"file": str(path), "ok": False}

    lengths = []
    gc = 0
    total = 0
    seq = []

    with path.open("r", encoding="utf-8", errors="replace") as f:
        for line in f:
            if line.startswith(">"):
                if seq:
                    s = "".join(seq).upper()
                    L = len(s)
                    lengths.append(L)
                    total += L
                    gc += s.count("G") + s.count("C")
                    seq = []
            else:
                seq.append(line.strip())
        if seq:
            s = "".join(seq).upper()
            L = len(s)
            lengths.append(L)
            total += L
            gc += s.count("G") + s.count("C")

    if total == 0 or not lengths:
        return {"file": str(path), "ok": False}

    lengths_sorted = sorted(lengths, reverse=True)
    half = total / 2.0
    csum = 0
    n50 = 0
    l50 = 0
    for i, L in enumerate(lengths_sorted, start=1):
        csum += L
        if csum >= half:
            n50 = L
            l50 = i
            break

    n = len(lengths)
    lengths_sorted2 = sorted(lengths)
    cum = 0
    for i, x in enumerate(lengths_sorted2, start=1):
        cum += i * x
    gini = (2 * cum) / (n * total) - (n + 1) / n

    return {
        "file": path.name,
        "ok": True,
        "contigs": n,
        "total_bp": total,
        "largest_bp": max(lengths),
        "gc_pct": (gc / total) * 100.0,
        "n50_bp": n50,
        "l50": l50,
        "gini_len": gini,
    }

def repeat_header_row(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_normal_style(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)

def set_cell(cell, text, bold=False, center=False):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text))
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT

def add_table(doc: Document, title: str, cols: list[str], rows: list[list], note: str | None = None):
    cap = doc.add_paragraph(title)
    cap.runs[0].bold = True

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, c in enumerate(cols):
        set_cell(hdr[j], c, bold=True, center=True)
    repeat_header_row(table.rows[0])

    for r in rows:
        cells = table.add_row().cells
        for j, v in enumerate(r):
            set_cell(cells[j], v, bold=False, center=(j != 0))

    if note:
        p = doc.add_paragraph(note)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph("")
    return table

def write_tsv(path: Path, cols: list[str], rows: list[list]):
    with path.open("w", newline="") as f:
        w = csv.writer(f, delimiter="\t")
        w.writerow(cols)
        w.writerows(rows)

def fmt_pct(x):
    try: return f"{float(x):.2f}"
    except Exception: return ""

def fmt_int(x):
    try: return f"{int(x)}"
    except Exception: return ""

def fmt_bp(x):
    try: return f"{int(x):,}"
    except Exception: return ""

# ----------------------------
# Table builders
# ----------------------------
def build_table1_kraken():
    dist_r = kraken_genus_distribution(parse_kraken_report(K_READS))
    dist_c = kraken_genus_distribution(parse_kraken_report(K_CONTIG))
    dist_m = kraken_genus_distribution(parse_kraken_report(K_MIN5))

    # union top genera
    def top_keys(dist, k=10):
        items = [(v, g) for g, v in dist.items() if g != "Unclassified"]
        items.sort(reverse=True)
        return [g for _, g in items[:k]]

    keys = set()
    for d in (dist_r, dist_c, dist_m):
        keys.update(top_keys(d, 10))
    keys = list(keys)
    keys.sort(key=lambda g: dist_r.get(g, 0.0), reverse=True)
    keys = keys[:12]

    rows = []
    if ("Unclassified" in dist_r) or ("Unclassified" in dist_c) or ("Unclassified" in dist_m):
        rows.append(["Unclassified", fmt_pct(dist_r.get("Unclassified","")),
                     fmt_pct(dist_c.get("Unclassified","")), fmt_pct(dist_m.get("Unclassified",""))])

    for g in keys:
        rows.append([g, fmt_pct(dist_r.get(g,"")), fmt_pct(dist_c.get(g,"")), fmt_pct(dist_m.get(g,""))])

    Hr, Sr = diversity_from_pct(dist_r)
    Hc, Sc = diversity_from_pct(dist_c)
    Hm, Sm = diversity_from_pct(dist_m) if dist_m else (0.0, 0.0)

    note = (f"Note: Kraken2 genus composition. Diversity: Shannon H (reads={Hr:.2f}, contigs={Hc:.2f}, ≥5kb={Hm:.2f}); "
            f"Simpson (1−D) (reads={Sr:.2f}, contigs={Sc:.2f}, ≥5kb={Sm:.2f}).")
    return rows, note

def build_table2_assembly():
    s_all = fasta_stats(FA_ALL)
    s_min = fasta_stats(FA_MIN5)
    s_bac = fasta_stats(FA_BAC)

    cols = ["Metric", "All contigs", "≥5 kb contigs", "Bacillus-only"]
    rows = [
        ["Contigs (#)", fmt_int(s_all.get("contigs","")), fmt_int(s_min.get("contigs","")), fmt_int(s_bac.get("contigs",""))],
        ["Total length (bp)", fmt_bp(s_all.get("total_bp","")), fmt_bp(s_min.get("total_bp","")), fmt_bp(s_bac.get("total_bp",""))],
        ["Largest contig (bp)", fmt_bp(s_all.get("largest_bp","")), fmt_bp(s_min.get("largest_bp","")), fmt_bp(s_bac.get("largest_bp",""))],
        ["GC (%)", fmt_pct(s_all.get("gc_pct","")), fmt_pct(s_min.get("gc_pct","")), fmt_pct(s_bac.get("gc_pct",""))],
        ["N50 (bp)", fmt_bp(s_all.get("n50_bp","")), fmt_bp(s_min.get("n50_bp","")), fmt_bp(s_bac.get("n50_bp",""))],
        ["L50", fmt_int(s_all.get("l50","")), fmt_int(s_min.get("l50","")), fmt_int(s_bac.get("l50",""))],
        ["Gini fragmentation (lengths)", fmt_pct(s_all.get("gini_len","")), fmt_pct(s_min.get("gini_len","")), fmt_pct(s_bac.get("gini_len",""))],
    ]
    note = "Note: Metrics computed directly from FASTA. Gini quantifies fragmentation (higher = more unequal contig length distribution)."
    return cols, rows, note

def build_table3_concordance():
    # detect clean tables first (preferred)
    amrf = try_paths([TBL_DIR / "I-9.amrfinder.clean.tsv", TBL_BAC / "I-9.bacillus_only.amrfinder.clean.tsv", I9 / "amr" / "I-9.amrfinder.tsv"])
    card = try_paths([TBL_DIR / "I-9.abricate.card.clean.tsv", I9 / "amr" / "I-9.abricate.card.tab"])
    ncbi = try_paths([TBL_DIR / "I-9.abricate.ncbi.clean.tsv", I9 / "amr" / "I-9.abricate.ncbi.tab"])
    vfdb = try_paths([TBL_DIR / "I-9.abricate.vfdb.clean.tsv", I9 / "amr" / "I-9.abricate.vfdb.tab"])
    plas = try_paths([TBL_DIR / "I-9.abricate.plasmidfinder.clean.tsv", I9 / "amr" / "I-9.abricate.plasmidfinder.tab"])

    gene = defaultdict(lambda: {
        "Gene": "", "Class": "",
        "AMRFinder": "0", "CARD": "0", "NCBI": "0", "VFDB": "0", "PlasmidFinder": "0",
        "BestIdentity": "", "BestCoverage": ""
    })

    def ingest_amrf(path: Path | None):
        if not path: return
        try:
            data = read_tsv(path)
            for r in data:
                g = (r.get("Gene symbol") or r.get("Gene") or r.get("Element") or "").strip()
                if not g: continue
                cls = (r.get("Class") or r.get("Subclass") or "").strip()
                gene[g]["Gene"] = g
                if cls and not gene[g]["Class"]:
                    gene[g]["Class"] = cls
                gene[g]["AMRFinder"] = "1"
        except Exception:
            return

    def ingest_abricate(path: Path | None, flag: str):
        if not path: return
        try:
            data = read_tsv(path)
            for r in data:
                g = (r.get("GENE") or r.get("Gene") or r.get("gene") or r.get("ID") or "").strip()
                if not g: continue
                gene[g]["Gene"] = g
                gene[g][flag] = "1"
                ident = (r.get("%IDENTITY") or r.get("IDENTITY") or r.get("identity") or "")
                cov   = (r.get("%COVERAGE") or r.get("COVERAGE") or r.get("coverage") or "")
                try:
                    ident_f = float(str(ident).replace("%",""))
                except Exception:
                    ident_f = None
                try:
                    cov_f = float(str(cov).replace("%",""))
                except Exception:
                    cov_f = None
                if ident_f is not None:
                    cur = gene[g]["BestIdentity"]
                    try:
                        curf = float(cur) if cur else -1.0
                    except Exception:
                        curf = -1.0
                    if ident_f > curf:
                        gene[g]["BestIdentity"] = f"{ident_f:.2f}"
                        if cov_f is not None:
                            gene[g]["BestCoverage"] = f"{cov_f:.2f}"
        except Exception:
            return

    ingest_amrf(amrf)
    ingest_abricate(card, "CARD")
    ingest_abricate(ncbi, "NCBI")
    ingest_abricate(vfdb, "VFDB")
    ingest_abricate(plas, "PlasmidFinder")

    rows = list(gene.values())
    def score(r): return sum(int(r[k]) for k in ["AMRFinder","CARD","NCBI","VFDB","PlasmidFinder"])
    rows.sort(key=lambda r: (score(r), r["Gene"].lower()), reverse=True)
    rows = rows[:40]

    cols = ["Gene","Class","AMRFinder","CARD","NCBI","VFDB","PlasmidFinder","Best identity (%)","Best coverage (%)"]
    out = [[r["Gene"], r["Class"], r["AMRFinder"], r["CARD"], r["NCBI"], r["VFDB"], r["PlasmidFinder"], r["BestIdentity"], r["BestCoverage"]] for r in rows]
    note = "Note: Presence flags are conservative (1 if gene label observed). Best identity/coverage taken from Abricate when available."
    return cols, out, note

# ----------------------------
# Main
# ----------------------------
def main():
    doc = Document()
    set_normal_style(doc)

    title = doc.add_paragraph("Manuscript-ready tables — Bacillus WGS (I-9)")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    doc.add_paragraph(f"Auto-generated from: {I9}")
    doc.add_paragraph("")

    # Table 1
    t1_rows, t1_note = build_table1_kraken()
    add_table(doc,
        "Table 1. Kraken2 genus composition (reads vs contigs vs contigs ≥ 5 kb)",
        ["Genus","Reads (%)","Contigs (%)","Contigs ≥5kb (%)"],
        t1_rows,
        t1_note
    )
    write_tsv(OUTDIR / "Table1_Kraken2_Genus.tsv",
              ["Genus","Reads_pct","Contigs_pct","Contigs_ge5kb_pct"], t1_rows)

    # Table 2
    t2_cols, t2_rows, t2_note = build_table2_assembly()
    add_table(doc,
        "Table 2. Assembly statistics across filters (computed from FASTA)",
        t2_cols, t2_rows, t2_note
    )
    write_tsv(OUTDIR / "Table2_Assembly_Stats.tsv", t2_cols, t2_rows)

    # Table 3
    t3_cols, t3_rows, t3_note = build_table3_concordance()
    add_table(doc,
        "Table 3. Tool concordance (AMRFinder + Abricate databases)",
        t3_cols, t3_rows, t3_note
    )
    write_tsv(OUTDIR / "Table3_Tool_Concordance.tsv", t3_cols, t3_rows)

    out_docx = OUTDIR / "Manuscript_Tables_I9.docx"
    doc.save(out_docx)

    print("Wrote:", out_docx)
    print("Wrote:", OUTDIR / "Table1_Kraken2_Genus.tsv")
    print("Wrote:", OUTDIR / "Table2_Assembly_Stats.tsv")
    print("Wrote:", OUTDIR / "Table3_Tool_Concordance.tsv")

if __name__ == "__main__":
    main()
