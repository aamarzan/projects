import argparse
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_font(run, name="Times New Roman", size=10, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold

def main(ae_csv, mort_csv, out_docx):
    ae = pd.read_csv(ae_csv)
    mort = pd.read_csv(mort_csv)

    # Build mortality single row mapped to the same table schema
    mort_row = {
        "Outcome": "Mortality",
        "Components / symptoms included": "—",
        "Overall (n)": str(mort.loc[0, "Overall (n)"]) if "Overall (n)" in mort.columns else "",
        "Mixed OP n (%)": str(mort.loc[0, "Mixed OP"]),
        "Chlorpyrifos n (%)": str(mort.loc[0, "Chlorpyrifos"]),
        "Cypermethrin n (%)": str(mort.loc[0, "Cypermethrin"]),
        "Risk diff (Mixed–Chlor) % (95% CI)": "—",
        "OR (Mixed vs Chlor) (95% CI)": str(mort.loc[0, "OR (Mixed vs Chlor) (95% CI)"]),
        "Mortality rate 95% CI (Mixed vs Chlor)": str(mort.loc[0, "95% CI (Mortality rate) Mixed vs Chlor"]),
        "p (Fisher) MC": str(mort.loc[0, "p (Fisher) MC"]),
        "Risk diff (Mixed–Cyp) % (95% CI)": "—",
        "OR (Mixed vs Cyp) (95% CI)": str(mort.loc[0, "OR (Mixed vs Cyp) (95% CI)"]),
        "Mortality rate 95% CI (Mixed vs Cyp)": str(mort.loc[0, "95% CI (Mortality rate) Mixed vs Cyp"]),
        "p (Fisher) MY": str(mort.loc[0, "p (Fisher) MY"]),
    }

    # AE rows mapped
    ae_rows = []
    for _, r in ae.iterrows():
        ae_rows.append({
            "Outcome": str(r["AE Group"]),
            "Components / symptoms included": str(r["Symptoms included"]),
            "Overall (n)": "",
            "Mixed OP n (%)": str(r["Mixed OP n (%)"]),
            "Chlorpyrifos n (%)": str(r["Chlorpyrifos n (%)"]),
            "Cypermethrin n (%)": str(r["Cypermethrin n (%)"]),
            "Risk diff (Mixed–Chlor) % (95% CI)": str(r["Risk diff (Mixed−Chlor) % (95% CI)"]),
            "OR (Mixed vs Chlor) (95% CI)": str(r["OR (Mixed vs Chlor) (95% CI)"]),
            "Mortality rate 95% CI (Mixed vs Chlor)": "",
            "p (Fisher) MC": str(r["p (Fisher) MC"]),
            "Risk diff (Mixed–Cyp) % (95% CI)": str(r["Risk diff (Mixed−Cyp) % (95% CI)"]),
            "OR (Mixed vs Cyp) (95% CI)": str(r["OR (Mixed vs Cyp) (95% CI)"]),
            "Mortality rate 95% CI (Mixed vs Cyp)": "",
            "p (Fisher) MY": str(r["p (Fisher) MY"]),
        })

    headers = list(mort_row.keys())  # same order as mort_row dict

    doc = Document()

    # Landscape + margins
    sec = doc.sections[0]
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in (("left_margin", 0.5), ("right_margin", 0.5), ("top_margin", 0.5), ("bottom_margin", 0.5)):
        setattr(sec, m[0], Inches(m[1]))

    # Caption
    p = doc.add_paragraph()
    run = p.add_run("Table 2. Grouped adverse events and mortality by pesticide exposure group")
    set_font(run, bold=True, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = doc.add_paragraph("Data are presented as number/total (%) unless otherwise noted. P values are two-sided Fisher’s exact test.")
    set_font(p2.runs[0], size=10)

    # Single table with section rows
    ncols = len(headers)
    nrows = 1 + 1 + len(ae_rows) + 1 + 1  # header + AE label + AE + Mortality label + Mortality
    t = doc.add_table(rows=nrows, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = h
        shade_cell(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=9)

    # Section row helper
    def write_section(row_idx, title):
        c0 = t.cell(row_idx, 0)
        c0.text = title
        for run in c0.paragraphs[0].runs:
            set_font(run, bold=True, size=9)
        # clear other cells
        for j in range(1, ncols):
            t.cell(row_idx, j).text = ""

    # AE section
    r = 1
    write_section(r, "A. Adverse events (grouped)")
    r += 1
    for row in ae_rows:
        for j, h in enumerate(headers):
            t.cell(r, j).text = row[h]
            for run in t.cell(r, j).paragraphs[0].runs:
                set_font(run, size=9)
        r += 1

    # Mortality section
    write_section(r, "B. Mortality")
    r += 1
    for j, h in enumerate(headers):
        t.cell(r, j).text = mort_row[h]
        for run in t.cell(r, j).paragraphs[0].runs:
            set_font(run, size=9)

    # Footnote
    doc.add_paragraph(
        "Abbreviations: RD, risk difference; OR, odds ratio; CI, confidence interval; "
        "MC, Mixed OP vs Chlorpyrifos; MY, Mixed OP vs Cypermethrin."
    )

    doc.save(out_docx)
    print(f"Saved: {out_docx}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--ae", required=True, help="AE_Grouped CSV")
    ap.add_argument("--mort", required=True, help="Mortality CSV")
    ap.add_argument("--out", default="Table_2.docx", help="Output DOCX filename")
    args = ap.parse_args()
    main(args.ae, args.mort, args.out)
