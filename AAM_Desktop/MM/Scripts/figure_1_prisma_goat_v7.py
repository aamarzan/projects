#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
figure_1_prisma_goat_v7.py

v7 changes:
1) Box widths are parameterized (CLI knobs) so boxes can be made tighter.
2) Bottom 3 panels are centered as a group -> equal left/right margins, no cropping.
3) Extra save padding to protect rounded corners/shadows.

Run:
  python figure_1_prisma_goat_v7.py --write_template prisma_counts.json
  # edit prisma_counts.json
  python figure_1_prisma_goat_v7.py --counts prisma_counts.json --outdir figures --bottom_panel --docx "Multiple Myeloma.docx" --eps

Most useful tuning knobs:
  --main_w        width of left/main spine boxes
  --side_w        width of right-side boxes
  --col_gap       gap between main and side columns
  --xmax          internal canvas width (increase if you add lots more text)
  --margin        left/right safe margin inside canvas
  --bottom_ratios ratios for (Search snapshot, BD summary, Diagnostics) widths, e.g. "1.2,1.0,0.9"
  --bottom_gap    gap between bottom panels
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


# --------------------------- Helpers / Style ---------------------------

def best_font() -> str:
    preferred = ["Calibri", "Arial", "Helvetica", "Liberation Sans", "DejaVu Sans"]
    available = {f.name for f in mpl.font_manager.fontManager.ttflist}
    for f in preferred:
        if f in available:
            return f
    return "DejaVu Sans"


def set_rc():
    f = best_font()
    mpl.rcParams.update(
        {
            "font.family": f,
            "font.size": 10.4,
            "figure.titlesize": 16,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
        }
    )


def wrap_keep_newlines(text: str, width: int) -> str:
    """Wrap while preserving explicit newlines."""
    if text is None:
        return ""
    text = str(text).rstrip("\n")
    if not text.strip():
        return ""
    paras = text.splitlines()
    out: List[str] = []
    for p in paras:
        p = re.sub(r"\s+", " ", p.strip())
        if not p:
            out.append("")
            continue
        words = p.split(" ")
        cur: List[str] = []
        n = 0
        for w in words:
            add = len(w) + (1 if cur else 0)
            if n + add > width:
                out.append(" ".join(cur))
                cur = [w]
                n = len(w)
            else:
                cur.append(w)
                n += add
        if cur:
            out.append(" ".join(cur))
    return "\n".join(out)


def gradient_box(
    ax,
    x,
    y,
    w,
    h,
    c0,
    c1,
    edge="#1f1f1f",
    lw=1.25,
    r=0.16,
    z=2,
    shadow=True,
    shadow_dx=0.055,
    shadow_dy=-0.055,
    shadow_alpha=0.08,
):
    if shadow:
        sh = FancyBboxPatch(
            (x + shadow_dx, y + shadow_dy),
            w,
            h,
            boxstyle=f"round,pad=0.012,rounding_size={r}",
            linewidth=0,
            edgecolor="none",
            facecolor="#000000",
            alpha=shadow_alpha,
            zorder=z - 1,
        )
        ax.add_patch(sh)

    n = 256
    grad = np.linspace(0, 1, n).reshape(n, 1)
    cmap = LinearSegmentedColormap.from_list("g", [c0, c1])
    im = ax.imshow(
        grad,
        extent=[x, x + w, y, y + h],
        origin="lower",
        cmap=cmap,
        interpolation="bicubic",
        zorder=z,
    )

    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.012,rounding_size={r}",
        linewidth=lw,
        edgecolor=edge,
        facecolor="none",
        zorder=z + 1,
    )
    ax.add_patch(patch)
    im.set_clip_path(patch)
    return patch


def label(
    ax,
    x,
    y,
    s,
    ha="center",
    va="center",
    size=10.4,
    weight="regular",
    color="#111111",
    z=10,
):
    ax.text(
        x,
        y,
        s,
        ha=ha,
        va=va,
        fontsize=size,
        fontweight=weight,
        color=color,
        zorder=z,
        linespacing=1.18,
    )


def arrow(ax, p0, p1, color="#262626", lw=1.35, ms=12.5, z=6, connectionstyle=None):
    kw = {}
    if connectionstyle:
        kw["connectionstyle"] = connectionstyle
    a = FancyArrowPatch(
        p0,
        p1,
        arrowstyle="-|>",
        mutation_scale=ms,
        linewidth=lw,
        color=color,
        zorder=z,
        **kw,
    )
    ax.add_patch(a)
    return a


def connect_down(ax, upper_box, lower_box, color="#262626", lw=1.45, ms=12.5):
    """Clean downward arrow between vertically stacked boxes."""
    ux, uy, uw, uh = upper_box
    lx, ly, lw2, lh = lower_box
    x = ux + uw / 2
    upper_bottom = uy
    lower_top = ly + lh
    gap = upper_bottom - lower_top
    m = min(0.08, max(0.035, gap * 0.32)) if gap > 0 else 0.05
    y0 = upper_bottom - m
    y1 = lower_top + m
    if y0 <= y1:
        y0 = upper_bottom - 0.035
        y1 = lower_top + 0.035
    arrow(ax, (x, y0), (x, y1), color=color, lw=lw, ms=ms)


# --------------------------- Data ---------------------------

@dataclass
class PrismaCounts:
    db: int
    reg: int
    web: int
    org: int
    cite: int
    dup: int
    auto: int
    other_rm: int
    screened: int
    excluded_ta: int
    reports_sought: int
    not_retrieved: int
    assessed: int
    fulltext_excl: List[Tuple[str, int]]
    studies: int
    reports: int
    meta_studies: int
    search_date: Optional[str] = None
    databases_list: Optional[List[str]] = None
    limits: Optional[str] = None
    notes: Optional[str] = None

    @staticmethod
    def from_json(d: Dict[str, Any]) -> "PrismaCounts":
        ident = d.get("records_identified", {})
        other = d.get("records_identified_other_methods", {})
        rm = d.get("records_removed_before_screening", {})
        ft = d.get("reports_excluded", []) or []
        fulltext: List[Tuple[str, int]] = []
        for item in ft:
            r = str(item.get("reason", "")).strip()
            n = int(item.get("n", 0))
            if r:
                fulltext.append((r, n))

        meta = d.get("metadata", {})
        return PrismaCounts(
            db=int(ident.get("databases", 0)),
            reg=int(ident.get("registers", 0)),
            web=int(other.get("websites", 0)),
            org=int(other.get("organisations", 0)),
            cite=int(other.get("citation_searching", 0)),
            dup=int(rm.get("duplicates", 0)),
            auto=int(rm.get("automation", 0)),
            other_rm=int(rm.get("other", 0)),
            screened=int(d.get("records_screened", 0)),
            excluded_ta=int(d.get("records_excluded", 0)),
            reports_sought=int(d.get("reports_sought_for_retrieval", 0)),
            not_retrieved=int(d.get("reports_not_retrieved", 0)),
            assessed=int(d.get("reports_assessed_for_eligibility", 0)),
            fulltext_excl=fulltext,
            studies=int(d.get("studies_included_in_review", 0)),
            reports=int(d.get("reports_included_in_review", 0)),
            meta_studies=int(d.get("studies_included_in_meta_analysis", 0)),
            search_date=str(meta.get("last_search_date")) if meta.get("last_search_date") else None,
            databases_list=meta.get("databases"),
            limits=str(meta.get("limits")) if meta.get("limits") else None,
            notes=str(meta.get("notes")) if meta.get("notes") else None,
        )


def extract_table1_summary(docx_path: Path) -> Optional[Dict[str, Any]]:
    """Find first table that contains 'Method' and 'Sample size' (case-insensitive)."""
    try:
        from docx import Document
    except Exception:
        return None
    if not docx_path.exists():
        return None
    doc = Document(str(docx_path))
    if not doc.tables:
        return None

    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip()).lower()

    target = None
    idx_n = idx_m = None
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = [norm(c.text) for c in t.rows[0].cells]
        if "sample size" in header and "method" in header:
            target = t
            idx_n = header.index("sample size")
            idx_m = header.index("method")
            break
    if target is None:
        return None

    studies = 0
    n_sum = 0
    n_counted = 0
    designs: Dict[str, int] = {}

    for r in target.rows[1:]:
        studies += 1
        n_raw = (r.cells[idx_n].text or "").strip()
        m_raw = (r.cells[idx_m].text or "").strip()
        m_key = re.sub(r"\s+", " ", m_raw).strip() or "NR"
        designs[m_key] = designs.get(m_key, 0) + 1
        if re.search(r"\d", n_raw):
            try:
                n = int(re.sub(r"[^\d]", "", n_raw))
                if n > 0:
                    n_sum += n
                    n_counted += 1
            except Exception:
                pass

    return {"studies": studies, "total_n": n_sum if n_counted else None, "n_counted": n_counted, "designs": designs}


# --------------------------- Layout knobs ---------------------------

@dataclass
class Layout:
    xmax: float = 21.8
    y_max: float = 10.0
    margin: float = 1.35
    main_w: float = 8.15
    side_w: float = 8.10
    col_gap: float = 1.00


def draw_prisma(ax, c: PrismaCounts, layout: Layout, subtitle: Optional[str] = None):
    ax.set_axis_off()
    ax.set_xlim(0, layout.xmax)
    ax.set_ylim(0, layout.y_max)

    border = "#1f1f1f"
    arrowc = "#262626"

    g_id = ("#ffffff", "#dfeeff")
    g_sc = ("#ffffff", "#fff0cc")
    g_el = ("#ffffff", "#e2f7ee")
    g_in = ("#ffffff", "#ffe4ef")
    g_sd = ("#ffffff", "#f0f0f0")

    label(ax, layout.xmax / 2, 9.74, "PRISMA flow diagram", size=16, weight="bold")
    if subtitle:
        label(ax, layout.xmax / 2, 9.38, subtitle, size=10.8, color="#444444")

    xM = layout.margin
    wM = layout.main_w
    xS = xM + wM + layout.col_gap
    wS = layout.side_w

    # shift left if user makes it too wide
    right_edge = xS + wS
    if right_edge > layout.xmax - 0.45:
        shift = right_edge - (layout.xmax - 0.45)
        xM = max(0.6, xM - shift)
        xS = xM + wM + layout.col_gap

    h = 0.90
    h_tall = 1.05
    dy = 1.18
    y0 = 8.28

    b_ident  = (xM, y0,                 wM, h_tall)
    b_unique = (xM, y0 - 1 * dy,        wM, h)
    b_screen = (xM, y0 - 2 * dy,        wM, h)
    b_sought = (xM, y0 - 3 * dy,        wM, h)
    b_retr   = (xM, y0 - 4 * dy,        wM, h)
    b_assess = (xM, y0 - 5 * dy,        wM, h)
    b_incl   = (xM, y0 - 6 * dy,        wM, h)
    b_meta   = (xM, y0 - 7 * dy,        wM, h)

    b_removed = (xS, y0 - 0.42 * dy,     wS, h_tall)
    b_excl_ta = (xS, b_screen[1],        wS, h)
    b_notret  = (xS, b_sought[1],        wS, h)
    b_excl_ft = (xS, b_assess[1] - 0.62, wS, 2.10)

    gradient_box(ax, *b_ident,  *g_id, edge=border, lw=1.32)
    gradient_box(ax, *b_unique, *g_id, edge=border, lw=1.22)
    gradient_box(ax, *b_screen, *g_sc, edge=border, lw=1.32)
    gradient_box(ax, *b_sought, *g_sc, edge=border, lw=1.22)
    gradient_box(ax, *b_retr,   *g_el, edge=border, lw=1.22)
    gradient_box(ax, *b_assess, *g_el, edge=border, lw=1.32)
    gradient_box(ax, *b_incl,   *g_in, edge=border, lw=1.32)
    gradient_box(ax, *b_meta,   *g_in, edge=border, lw=1.22)

    gradient_box(ax, *b_removed, *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_excl_ta, *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_notret,  *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_excl_ft, *g_sd, edge=border, lw=1.05)

    label(ax, xM - 0.55, b_ident[1] + b_ident[3] / 2, "Identification", ha="right", size=10.6, weight="bold", color="#2b4d8a")
    label(ax, xM - 0.55, b_screen[1] + h / 2,         "Screening",      ha="right", size=10.6, weight="bold", color="#8a5a00")
    label(ax, xM - 0.55, b_assess[1] + h / 2,         "Eligibility",    ha="right", size=10.6, weight="bold", color="#1f6b3b")
    label(ax, xM - 0.55, b_meta[1] + h / 2,           "Included",       ha="right", size=10.6, weight="bold", color="#7a1f4c")

    ident_txt = (
        "Records identified from\n"
        f"databases (n = {c.db}), registers (n = {c.reg}),\n"
        f"websites (n = {c.web}), organisations (n = {c.org}),\n"
        f"citation searching (n = {c.cite})"
    )
    label(ax, b_ident[0] + wM / 2, b_ident[1] + b_ident[3] / 2, wrap_keep_newlines(ident_txt, 44), size=10.0, weight="bold")

    unique_n = max(0, (c.db + c.reg + c.web + c.org + c.cite) - (c.dup + c.auto + c.other_rm))
    label(ax, b_unique[0] + wM / 2, b_unique[1] + h / 2, f"Unique records after removals\n(n = {unique_n})", size=11.0, weight="bold")

    removed_txt = (
        "Records removed before screening\n"
        f"Duplicate records removed (n = {c.dup})\n"
        f"Marked as ineligible by automation (n = {c.auto})\n"
        f"Removed for other reasons (n = {c.other_rm})"
    )
    label(ax, b_removed[0] + wS / 2, b_removed[1] + b_removed[3] / 2, wrap_keep_newlines(removed_txt, 52), ha="center", va="center", size=9.0)

    label(ax, b_screen[0] + wM / 2, b_screen[1] + h / 2, f"Records screened\n(n = {c.screened})", size=11.2, weight="bold")
    label(ax, b_excl_ta[0] + wS / 2, b_excl_ta[1] + h / 2, f"Records excluded\n(n = {c.excluded_ta})", size=10.8, weight="bold")

    label(ax, b_sought[0] + wM / 2, b_sought[1] + h / 2, f"Reports sought for retrieval\n(n = {c.reports_sought})", size=10.9, weight="bold")
    label(ax, b_notret[0] + wS / 2, b_notret[1] + h / 2, f"Reports not retrieved\n(n = {c.not_retrieved})", size=10.8, weight="bold")

    retrieved_n = max(0, c.reports_sought - c.not_retrieved)
    label(ax, b_retr[0] + wM / 2, b_retr[1] + h / 2, f"Reports retrieved\n(n = {retrieved_n})", size=10.9, weight="bold")

    label(ax, b_assess[0] + wM / 2, b_assess[1] + h / 2, f"Reports assessed for eligibility\n(n = {c.assessed})", size=10.9, weight="bold")

    if c.fulltext_excl:
        lines = [f"• {r} (n = {n})" for r, n in c.fulltext_excl]
        ft_txt = "Reports excluded (with reasons)\n" + "\n".join(lines)
    else:
        ft_txt = "Reports excluded (with reasons)\n• Add full-text exclusion reasons in JSON"
    label(ax, b_excl_ft[0] + wS / 2, b_excl_ft[1] + b_excl_ft[3] / 2, wrap_keep_newlines(ft_txt, 56), ha="center", va="center", size=9.0)

    label(ax, b_incl[0] + wM / 2, b_incl[1] + h / 2, f"Studies included in review (qualitative)\n(n = {c.studies})\nReports included\n(n = {c.reports})", size=10.2, weight="bold")
    label(ax, b_meta[0] + wM / 2, b_meta[1] + h / 2, f"Studies included in meta-analysis (quantitative)\n(n = {c.meta_studies})", size=10.5, weight="bold")

    connect_down(ax, b_ident, b_unique, color=arrowc)
    connect_down(ax, b_unique, b_screen, color=arrowc)
    connect_down(ax, b_screen, b_sought, color=arrowc)
    connect_down(ax, b_sought, b_retr, color=arrowc)
    connect_down(ax, b_retr, b_assess, color=arrowc)
    connect_down(ax, b_assess, b_incl, color=arrowc)
    connect_down(ax, b_incl, b_meta, color=arrowc)

    elbow = "angle3,angleA=0,angleB=90"
    arrow(ax, (xM + wM + 0.12, b_ident[1] + b_ident[3] / 2), (xS - 0.10, b_removed[1] + b_removed[3] / 2), color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_screen[1] + h / 2), (xS - 0.10, b_excl_ta[1] + h / 2), color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_sought[1] + h / 2), (xS - 0.10, b_notret[1] + h / 2), color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_assess[1] + h / 2), (xS - 0.10, b_excl_ft[1] + b_excl_ft[3] / 2), color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)


def parse_ratios(s: str) -> Tuple[float, float, float]:
    parts = [p.strip() for p in (s or "").split(",") if p.strip()]
    if len(parts) != 3:
        return (1.10, 1.00, 0.95)
    try:
        r = [float(x) for x in parts]
        if min(r) <= 0:
            return (1.10, 1.00, 0.95)
        return (r[0], r[1], r[2])
    except Exception:
        return (1.10, 1.00, 0.95)


def draw_bottom_panel(ax, c: PrismaCounts, bd_summary: Optional[Dict[str, Any]], layout: Layout,
                      bottom_gap: float = 0.85, bottom_ratios: Tuple[float, float, float] = (1.10, 1.00, 0.95)):
    ax.set_axis_off()
    ax.set_xlim(0, layout.xmax)
    ax.set_ylim(0, 2.85)
    border = "#1f1f1f"

    y = 0.42
    h = 2.05
    gap = bottom_gap
    r1, r2, r3 = bottom_ratios

    usable = (layout.xmax - 2 * layout.margin) - 2 * gap
    total_r = r1 + r2 + r3
    w1 = usable * (r1 / total_r)
    w2 = usable * (r2 / total_r)
    w3 = usable * (r3 / total_r)

    total_w = w1 + w2 + w3 + 2 * gap
    x0 = (layout.xmax - total_w) / 2.0  # <-- centers the 3 boxes (equal blank space)

    gradient_box(ax, x0, y, w1, h, "#ffffff", "#eef5ff", edge=border, lw=1.0, r=0.16, shadow=True)
    label(ax, x0 + 0.35, y + h - 0.30, "Search snapshot", ha="left", size=11.0, weight="bold")

    lines = []
    if c.search_date: lines.append(f"Last search date: {c.search_date}")
    if c.databases_list: lines.append("Databases: " + ", ".join(c.databases_list))
    if c.limits: lines.append("Limits: " + c.limits)
    if c.notes: lines.append("Notes: " + c.notes)
    if not lines:
        lines = [
            "Last search date: edit metadata.last_search_date in JSON",
            "Databases: edit metadata.databases in JSON",
            "Limits: optional",
        ]
    label(ax, x0 + 0.35, y + h - 0.65, wrap_keep_newlines("\n".join(lines), 70), ha="left", va="top", size=8.8, color="#222222")

    x1 = x0 + w1 + gap
    gradient_box(ax, x1, y, w2, h, "#ffffff", "#eefcf2", edge=border, lw=1.0, r=0.16, shadow=True)
    label(ax, x1 + 0.35, y + h - 0.30, "Bangladesh evidence summary", ha="left", size=11.0, weight="bold")

    if bd_summary:
        studies = bd_summary.get("studies")
        total_n = bd_summary.get("total_n")
        n_counted = bd_summary.get("n_counted")
        designs = bd_summary.get("designs", {})
        top = f"Extracted Table 1 studies: {studies}"
        if total_n is not None and n_counted:
            top += f"\nSum sample size (extractable): {total_n} (from {n_counted} studies)"
        design_str = "; ".join([f"{k}: {v}" for k, v in designs.items()]) if designs else "NR"
        body = top + "\nDesigns: " + design_str
        label(ax, x1 + 0.35, y + h - 0.65, wrap_keep_newlines(body, 62), ha="left", va="top", size=8.7, color="#222222")
    else:
        label(ax, x1 + 0.35, y + h - 0.65,
              "BD summary not available.\nRun with --docx and ensure the table contains\ncolumns: Method, Sample size.",
              ha="left", va="top", size=8.7, color="#444444")

    x2 = x1 + w2 + gap
    gradient_box(ax, x2, y, w3, h, "#ffffff", "#fff0f4", edge=border, lw=1.0, r=0.16, shadow=True)
    label(ax, x2 + 0.35, y + h - 0.30, "Diagnostics", ha="left", size=11.0, weight="bold")
    diag = (
        "When JSON is final, enable validation:\n"
        "• arithmetic consistency\n"
        "• auto-highlight inconsistencies (red flags)"
    )
    label(ax, x2 + 0.35, y + h - 0.65, wrap_keep_newlines(diag, 60), ha="left", va="top", size=8.7, color="#333333")


def write_template(path: Path):
    tpl = {
        "records_identified": {"databases": 0, "registers": 0},
        "records_identified_other_methods": {"websites": 0, "organisations": 0, "citation_searching": 0},
        "records_removed_before_screening": {"duplicates": 0, "automation": 0, "other": 0},
        "records_screened": 0,
        "records_excluded": 0,
        "reports_sought_for_retrieval": 0,
        "reports_not_retrieved": 0,
        "reports_assessed_for_eligibility": 0,
        "reports_excluded": [
            {"reason": "Not Bangladesh population", "n": 0},
            {"reason": "Wrong outcome or non-MM plasma cell disorder", "n": 0},
            {"reason": "Case report or non-eligible design", "n": 0},
            {"reason": "Insufficient extractable data", "n": 0},
        ],
        "studies_included_in_review": 0,
        "reports_included_in_review": 0,
        "studies_included_in_meta_analysis": 0,
        "metadata": {
            "last_search_date": "DD-MMM-YYYY",
            "databases": ["PubMed", "Scopus", "Web of Science"],
            "limits": "Language: English; Humans; Bangladesh; up to Nov 2025",
            "notes": "If citation chasing was used after de-duplication, document it here.",
        },
    }
    path.write_text(json.dumps(tpl, indent=2), encoding="utf-8")


def main():
    set_rc()
    ap = argparse.ArgumentParser()
    ap.add_argument("--counts", type=str, default=None)
    ap.add_argument("--write_template", type=str, default=None)
    ap.add_argument("--docx", type=str, default=None)
    ap.add_argument("--bottom_panel", action="store_true")
    ap.add_argument("--outdir", type=str, default="figures")
    ap.add_argument("--basename", type=str, default="Figure_1_PRISMA_GOAT_v7")
    ap.add_argument("--dpi", type=int, default=600)
    ap.add_argument("--eps", action="store_true")
    ap.add_argument("--subtitle", type=str, default=None)

    ap.add_argument("--xmax", type=float, default=21.8)
    ap.add_argument("--margin", type=float, default=1.35)
    ap.add_argument("--main_w", type=float, default=8.15)
    ap.add_argument("--side_w", type=float, default=8.10)
    ap.add_argument("--col_gap", type=float, default=1.00)

    ap.add_argument("--bottom_gap", type=float, default=0.85)
    ap.add_argument("--bottom_ratios", type=str, default="1.10,1.00,0.95")

    args = ap.parse_args()

    if args.write_template:
        p = Path(args.write_template)
        p.parent.mkdir(parents=True, exist_ok=True)
        write_template(p)
        print(f"Wrote template to: {p}")
        return

    if not args.counts:
        raise SystemExit("ERROR: provide --counts prisma_counts.json (or create one with --write_template).")

    counts_path = Path(args.counts)
    if not counts_path.exists():
        raise SystemExit(f"ERROR: counts file not found: {counts_path}")

    d = json.loads(counts_path.read_text(encoding="utf-8"))
    c = PrismaCounts.from_json(d)

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    bd_summary = None
    if args.bottom_panel and args.docx:
        bd_summary = extract_table1_summary(Path(args.docx))

    layout = Layout(
        xmax=float(args.xmax),
        margin=float(args.margin),
        main_w=float(args.main_w),
        side_w=float(args.side_w),
        col_gap=float(args.col_gap),
    )
    ratios = parse_ratios(args.bottom_ratios)

    if args.bottom_panel:
        fig = plt.figure(figsize=(18.4, 10.1))
        gs = fig.add_gridspec(2, 1, height_ratios=[3.85, 1.15], hspace=0.02)
        ax0 = fig.add_subplot(gs[0, 0])
        ax1 = fig.add_subplot(gs[1, 0])
        draw_prisma(ax0, c, layout=layout, subtitle=args.subtitle)
        draw_bottom_panel(ax1, c, bd_summary, layout=layout, bottom_gap=float(args.bottom_gap), bottom_ratios=ratios)
    else:
        fig = plt.figure(figsize=(18.4, 8.0))
        ax0 = fig.add_subplot(1, 1, 1)
        draw_prisma(ax0, c, layout=layout, subtitle=args.subtitle)

    base = outdir / args.basename
    save_kws = dict(bbox_inches="tight", facecolor="white", pad_inches=0.20)
    fig.savefig(str(base) + ".png", dpi=args.dpi, **save_kws)
    fig.savefig(str(base) + ".pdf", **save_kws)
    fig.savefig(str(base) + ".svg", **save_kws)
    if args.eps:
        fig.savefig(str(base) + ".eps", **save_kws)
    plt.close(fig)

    print(f"Saved: {base}.png/.pdf/.svg" + ("/.eps" if args.eps else ""))


if __name__ == "__main__":
    main()
