import os
import sys
import pytesseract
from pdf2docx import Converter
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
import subprocess
import io

# ===== USER CONFIGURATION =====
TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_BIN_DIR = r'C:\poppler-24.08.0\Library\bin'
INPUT_FOLDER = r'C:\Users\User\OneDrive\2. Personal\AAM\Desktop\Image Conversion\Input Folder'
OUTPUT_FOLDER = r'C:\Users\User\OneDrive\2. Personal\AAM\Desktop\Image Conversion\Output Folder'

# ===== SYSTEM SETUP =====
def verify_poppler():
    """Ensure Poppler binaries are accessible"""
    pdftotext_path = os.path.join(POPPLER_BIN_DIR, 'pdftotext.exe')
    if not os.path.exists(pdftotext_path):
        print(f"❌ Critical: pdftotext.exe not found at {pdftotext_path}")
        print("Please verify:")
        print("1. Poppler is downloaded from https://github.com/oschwartz10612/poppler-windows/releases")
        print("2. The folder structure matches: C:\\poppler-24.08.0\\Library\\bin\\pdftotext.exe")
        return False
    
    # Add Poppler to temporary PATH for this process
    os.environ['PATH'] = f"{POPPLER_BIN_DIR}{os.pathsep}{os.environ['PATH']}"
    return True

def verify_tesseract():
    """Ensure Tesseract is accessible"""
    if not os.path.exists(TESSERACT_PATH):
        print(f"❌ Critical: tesseract.exe not found at {TESSERACT_PATH}")
        print("Download from: https://github.com/UB-Mannheim/tesseract/wiki")
        return False
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    return True

# ===== CONVERSION FUNCTIONS =====
def extract_text_with_position(img):
    # Current denoising (h=10 is medium strength)
    denoised = cv2.fastNlMeansDenoising(gray, h=10)  # ← Adjust this value
    
def extract_text_with_position(img):
    data = pytesseract.image_to_data(
        denoised,
        config='--psm 6 -c preserve_interword_spaces=1',  # ← Change to --psm 11
        output_type=pytesseract.Output.DICT
    )

def is_scanned_pdf(pdf_path):
    """Check if PDF contains selectable text"""
    try:
        result = subprocess.run(
            ['pdftotext', pdf_path, '-'],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        return len(result.stdout.strip()) < 50
    except Exception as e:
        print(f"⚠️ Warning: pdftotext check failed ({e}), assuming scanned PDF")
        return True

def convert_pdf(pdf_path, docx_path):
    """Convert PDF to DOCX with OCR fallback"""
    try:
        if is_scanned_pdf(pdf_path):
            print("🔍 Scanned PDF detected - using OCR")
            images = convert_from_path(
                pdf_path,
                dpi=600,
                poppler_path=POPPLER_BIN_DIR,
                grayscale=True
            )
            
            doc = Document()
            for img in images:
                text = pytesseract.image_to_string(img, config='--psm 6')
                doc.add_paragraph(text)
                
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                doc.add_picture(img_byte_arr, width=Inches(6.0))
            
            doc.save(docx_path)
        else:
            print("📝 Text-based PDF detected - direct conversion")
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
        return True
    except Exception as e:
        print(f"❌ Conversion failed: {str(e)}")
        return False

# ===== MAIN EXECUTION =====
if __name__ == "__main__":
    print("\n=== PDF to DOCX Converter ===")
    print(f"Input: {INPUT_FOLDER}")
    print(f"Output: {OUTPUT_FOLDER}")
    
    if not all([verify_poppler(), verify_tesseract()]):
        sys.exit(1)
    
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    for filename in os.listdir(INPUT_FOLDER):
        if filename.lower().endswith('.pdf'):
            input_path = os.path.join(INPUT_FOLDER, filename)
            output_name = f"{os.path.splitext(filename)[0]}_converted.docx"
            output_path = os.path.join(OUTPUT_FOLDER, output_name)
            
            print(f"\n📄 Processing: {filename}")
            if convert_pdf(input_path, output_path):
                print(f"✅ Success: {output_name}")
            else:
                print(f"❌ Failed: {filename}")
    
    print("\nConversion complete!")