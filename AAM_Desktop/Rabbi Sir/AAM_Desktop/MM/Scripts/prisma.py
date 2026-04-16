#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Figure_1_PRISMA.py

v13 (text clarity fixes):
- FIX: Newlines are real newlines (no literal '\\n' on the figure).
- NEW: 3-column layout option:
    Column 1: Main PRISMA spine (left flow boxes)
    Column 2: Standard PRISMA side boxes (4 boxes)
    Column 3: Supplementary panels (Search snapshot, Bangladesh evidence summary, Diagnostics)
  -> This removes the bottom row panels, saving vertical space and enabling larger, clearer text.
- NEW: Separate, in-script controllers for:
    (A) left flow boxes, (B) right PRISMA boxes, (C) supplementary panels.

Run (same as before):
  python Figure_1_PRISMA.py --counts prisma_counts.json --outdir figures --docx "Multiple Myeloma.docx"
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


# ============================ USER CONTROLLERS ============================
# 1) WIDTH CONTROLS (data-units)
#    Tip: if boxes look too wide, reduce widths and/or gaps. If anything crops, keep AUTO_XMAX=True.
AUTO_XMAX = True
MARGIN_X = 1.20
BOX_TEXT_SIZE = 9.6   # <- change only this to control ALL box text size
FORCE_GLOBAL_BOX_TEXT_SIZE = True

FLOW_W = 7.30           # left spine box width
SIDE_W = 7.10           # right PRISMA side-box width (4 boxes)
PANEL_W = 6.60          # rightmost panel width (3 boxes)

GAP_FLOW_TO_SIDE = 0.90
GAP_SIDE_TO_PANEL = 0.90

# 2) HEIGHT + GAP CONTROLS (data-units)
# Left spine boxes:
FLOW_H = 0.95
FLOW_H_TALL = 1.62      # top identification box
FLOW_V_GAP = 0.34       # vertical gap between left spine boxes

# Right PRISMA side boxes:
SIDE_H = 0.95
SIDE_H_TALL = 1.22      # top-right "removed" box height
SIDE_FT_H = 2.25        # "Reports excluded (with reasons)" box height

# Supplementary panels (rightmost column):
PANEL_GAP = 0.32
PANEL_H = 3.68          # each of the 3 panels height (stacked top->bottom)

# 3) TYPOGRAPHY
TITLE_SIZE = 16
BASE_FONT = 10.6
TITLE_WEIGHT = "bold"

# Box text defaults (label_in_box will auto-wrap + auto-shrink as needed)
FLOW_MAX = 11.3
FLOW_MIN = 8.4
SIDE_MAX = 10.8
SIDE_MIN = 8.0
PANEL_BODY_MAX = 9.2
PANEL_BODY_MIN = 7.4

# 4) VERTICAL CANVAS MARGINS
TOP_SPACE = 1.10        # reserved space for title above top boxes
BOTTOM_MARGIN = 0.55    # white space under lowest content
# ========================================================================


# --------------------------- Helpers / Style ---------------------------

def best_font() -> str:
    preferred = ["Calibri", "Arial", "Helvetica", "Liberation Sans", "DejaVu Sans"]
    available = {f.name for f in mpl.font_manager.fontManager.ttflist}
    for f in preferred:
        if f in available:
            return f
    return "DejaVu Sans"


def set_rc():
    f = best_font()
    mpl.rcParams.update(
        {
            "font.family": f,
            "font.size": BASE_FONT,
            "figure.titlesize": TITLE_SIZE,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
        }
    )


def wrap_keep_newlines(text: str, width: int) -> str:
    """Wrap while preserving explicit newlines."""
    if text is None:
        return ""
    text = str(text).rstrip("\n")
    if not text.strip():
        return ""
    paras = text.splitlines()
    out: List[str] = []
    for p in paras:
        p = re.sub(r"\s+", " ", p.strip())
        if not p:
            out.append("")
            continue
        words = p.split(" ")
        cur: List[str] = []
        n = 0
        for w in words:
            add = len(w) + (1 if cur else 0)
            if n + add > width:
                out.append(" ".join(cur))
                cur = [w]
                n = len(w)
            else:
                cur.append(w)
                n += add
        if cur:
            out.append(" ".join(cur))
    return "\n".join(out)


def gradient_box(
    ax,
    x,
    y,
    w,
    h,
    c0,
    c1,
    edge="#1f1f1f",
    lw=1.25,
    r=0.16,
    z=2,
    shadow=True,
    shadow_dx=0.055,
    shadow_dy=-0.055,
    shadow_alpha=0.08,
):
    if shadow:
        sh = FancyBboxPatch(
            (x + shadow_dx, y + shadow_dy),
            w,
            h,
            boxstyle=f"round,pad=0.012,rounding_size={r}",
            linewidth=0,
            edgecolor="none",
            facecolor="#000000",
            alpha=shadow_alpha,
            zorder=z - 1,
        )
        ax.add_patch(sh)

    n = 256
    grad = np.linspace(0, 1, n).reshape(n, 1)
    cmap = LinearSegmentedColormap.from_list("g", [c0, c1])
    im = ax.imshow(
        grad,
        extent=[x, x + w, y, y + h],
        origin="lower",
        cmap=cmap,
        interpolation="bicubic",
        zorder=z,
    )

    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0.012,rounding_size={r}",
        linewidth=lw,
        edgecolor=edge,
        facecolor="none",
        zorder=z + 1,
    )
    ax.add_patch(patch)
    im.set_clip_path(patch)
    return patch


def label(ax, x, y, s, ha="center", va="center", size=BASE_FONT, weight="regular", color="#111111", z=10):
    ax.text(
        x,
        y,
        s,
        ha=ha,
        va=va,
        fontsize=size,
        fontweight=weight,
        color=color,
        zorder=z,
        linespacing=1.18,
    )


def arrow(ax, p0, p1, color="#262626", lw=1.35, ms=12.5, z=6, connectionstyle=None):
    kw = {}
    if connectionstyle:
        kw["connectionstyle"] = connectionstyle
    a = FancyArrowPatch(
        p0,
        p1,
        arrowstyle="-|>",
        mutation_scale=ms,
        linewidth=lw,
        color=color,
        zorder=z,
        **kw,
    )
    ax.add_patch(a)
    return a


def connect_down(ax, upper_box, lower_box, color="#262626", lw=1.45, ms=12.5):
    """Clean downward arrow between vertically stacked boxes."""
    ux, uy, uw, uh = upper_box
    lx, ly, lw2, lh = lower_box
    x = ux + uw / 2
    upper_bottom = uy
    lower_top = ly + lh
    gap = upper_bottom - lower_top
    m = min(0.08, max(0.035, gap * 0.32)) if gap > 0 else 0.05
    y0 = upper_bottom - m
    y1 = lower_top + m
    if y0 <= y1:
        y0 = upper_bottom - 0.035
        y1 = lower_top + 0.035
    arrow(ax, (x, y0), (x, y1), color=color, lw=lw, ms=ms)


# --------------------------- Box-safe text helpers ---------------------------

def _wrap_width_for_box(box_w: float) -> int:
    """Heuristic: convert box width (data units) to approx characters/line."""
    return int(max(22, min(92, round(box_w * 7.2))))


def label_in_box(
    ax,
    box,
    text: str,
    *,
    ha: str = "center",
    va: str = "center",
    size: float = BOX_TEXT_SIZE,
    max_size: float | None = None,   # accepted for backward-compat (ignored if FORCE_GLOBAL_BOX_TEXT_SIZE)
    min_size: float | None = None,   # accepted for backward-compat (ignored if FORCE_GLOBAL_BOX_TEXT_SIZE)
    weight: str = "regular",
    color: str = "#111111",
    pad: float = 0.12,
    wrap: bool = True,
    justify: bool = False,
    clip_patch=None,
    z: int = 10,
):
    x, y, w, h = box
    s = "" if text is None else str(text)

    # enforce one global font size everywhere (your requirement)
    if FORCE_GLOBAL_BOX_TEXT_SIZE:
        size = BOX_TEXT_SIZE
    else:
        # if someone uses the old API and does not pass size, respect max_size
        if max_size is not None and (size == BOX_TEXT_SIZE):
            size = max_size

    if wrap:
        s = wrap_keep_newlines(s, _wrap_width_for_box(w))

    # anchor position
    if ha == "left":
        tx = x + pad
    elif ha == "right":
        tx = x + w - pad
    else:
        tx = x + w / 2

    if va == "top":
        ty = y + h - pad
    elif va == "bottom":
        ty = y + pad
    else:
        ty = y + h / 2

    # pseudo-justify (matplotlib has no true justify)
    if justify and ha == "left" and va == "top":
        ax.figure.canvas.draw()
        renderer = ax.figure.canvas.get_renderer()

        tmp = ax.text(0, 0, " ", fontsize=size, fontweight=weight, alpha=0)
        space_px = tmp.get_window_extent(renderer=renderer).width
        tmp.remove()

        max_px = (
            ax.transData.transform((x + w - pad, 0))[0]
            - ax.transData.transform((x + pad, 0))[0]
        )

        def text_px(line: str) -> float:
            t0 = ax.text(0, 0, line, fontsize=size, fontweight=weight, alpha=0)
            bb = t0.get_window_extent(renderer=renderer)
            t0.remove()
            return bb.width

        out_lines = []
        lines = s.splitlines()
        for i, line in enumerate(lines):
            words = line.split()
            if i == len(lines) - 1 or len(words) <= 2:
                out_lines.append(line)
                continue

            w0 = text_px(line)
            extra = max_px - w0
            if extra <= space_px:
                out_lines.append(line)
                continue

            gaps = len(words) - 1
            add_spaces = int(round(extra / space_px))
            base = add_spaces // gaps
            rem = add_spaces % gaps

            new = []
            for gi, wd in enumerate(words[:-1]):
                new.append(wd)
                new.append(" " * (1 + base + (1 if gi < rem else 0)))
            new.append(words[-1])
            out_lines.append("".join(new))

        s = "\n".join(out_lines)

    t = ax.text(
        tx, ty, s,
        ha=ha, va=va,
        fontsize=size,
        fontweight=weight,
        color=color,
        zorder=z,
        linespacing=1.18,
    )

    if clip_patch is not None:
        t.set_clip_path(clip_patch)
        t.set_clip_on(True)

    return t


def _zero_phrase(n: int, *, kind: str) -> str:
    if n != 0:
        return ""
    if kind == "searched":
        return "not searched"
    if kind == "used":
        return "not used"
    if kind == "performed":
        return "not performed"
    return "not applicable"


def _line(label0: str, n: int, *, zero_kind: str = "searched") -> str:
    zp = _zero_phrase(n, kind=zero_kind)
    if n == 0 and zp:
        return f"{label0} {zp} (n = 0)"
    return f"{label0} (n = {n})"


# --------------------------- Data ---------------------------

@dataclass
class PrismaCounts:
    db: int
    reg: int
    web: int
    org: int
    cite: int
    dup: int
    auto: int
    other_rm: int
    screened: int
    excluded_ta: int
    reports_sought: int
    not_retrieved: int
    assessed: int
    fulltext_excl: List[Tuple[str, int]]
    studies: int
    reports: int
    meta_studies: int
    search_date: Optional[str] = None
    databases_list: Optional[List[str]] = None
    limits: Optional[str] = None
    notes: Optional[str] = None

    @staticmethod
    def from_json(d: Dict[str, Any]) -> "PrismaCounts":
        ident = d.get("records_identified", {})
        other = d.get("records_identified_other_methods", {})
        rm = d.get("records_removed_before_screening", {})
        ft = d.get("reports_excluded", []) or []
        fulltext: List[Tuple[str, int]] = []
        for item in ft:
            r = str(item.get("reason", "")).strip()
            n = int(item.get("n", 0))
            if r:
                fulltext.append((r, n))
        meta = d.get("metadata", {})
        return PrismaCounts(
            db=int(ident.get("databases", 0)),
            reg=int(ident.get("registers", 0)),
            web=int(other.get("websites", 0)),
            org=int(other.get("organisations", 0)),
            cite=int(other.get("citation_searching", 0)),
            dup=int(rm.get("duplicates", 0)),
            auto=int(rm.get("automation", 0)),
            other_rm=int(rm.get("other", 0)),
            screened=int(d.get("records_screened", 0)),
            excluded_ta=int(d.get("records_excluded", 0)),
            reports_sought=int(d.get("reports_sought_for_retrieval", 0)),
            not_retrieved=int(d.get("reports_not_retrieved", 0)),
            assessed=int(d.get("reports_assessed_for_eligibility", 0)),
            fulltext_excl=fulltext,
            studies=int(d.get("studies_included_in_review", 0)),
            reports=int(d.get("reports_included_in_review", 0)),
            meta_studies=int(d.get("studies_included_in_meta_analysis", 0)),
            search_date=str(meta.get("last_search_date")) if meta.get("last_search_date") else None,
            databases_list=meta.get("databases"),
            limits=str(meta.get("limits")) if meta.get("limits") else None,
            notes=str(meta.get("notes")) if meta.get("notes") else None,
        )


def extract_table1_summary(docx_path: Path) -> Optional[Dict[str, Any]]:
    """Find first table that contains 'Method' and 'Sample size' (case-insensitive)."""
    try:
        from docx import Document
    except Exception:
        return None
    if not docx_path.exists():
        return None
    doc = Document(str(docx_path))
    if not doc.tables:
        return None

    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip()).lower()

    target = None
    idx_n = idx_m = None
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = [norm(c.text) for c in t.rows[0].cells]
        if "sample size" in header and "method" in header:
            target = t
            idx_n = header.index("sample size")
            idx_m = header.index("method")
            break
    if target is None:
        return None

    studies = 0
    n_sum = 0
    n_counted = 0
    designs: Dict[str, int] = {}

    for r in target.rows[1:]:
        studies += 1
        n_raw = (r.cells[idx_n].text or "").strip()
        m_raw = (r.cells[idx_m].text or "").strip()
        m_key = re.sub(r"\s+", " ", m_raw).strip() or "NR"
        designs[m_key] = designs.get(m_key, 0) + 1
        if re.search(r"\d", n_raw):
            try:
                n = int(re.sub(r"[^\d]", "", n_raw))
                if n > 0:
                    n_sum += n
                    n_counted += 1
            except Exception:
                pass

    return {"studies": studies, "total_n": n_sum if n_counted else None, "n_counted": n_counted, "designs": designs}


# --------------------------- Layout ---------------------------

@dataclass
class Layout:
    # horizontal
    margin_x: float = MARGIN_X
    flow_w: float = FLOW_W
    side_w: float = SIDE_W
    panel_w: float = PANEL_W
    gap_flow_side: float = GAP_FLOW_TO_SIDE
    gap_side_panel: float = GAP_SIDE_TO_PANEL
    xmax: float = 0.0  # auto-filled

    # vertical
    top_space: float = TOP_SPACE
    bottom_margin: float = BOTTOM_MARGIN

    flow_h: float = FLOW_H
    flow_h_tall: float = FLOW_H_TALL
    flow_v_gap: float = FLOW_V_GAP

    side_h: float = SIDE_H
    side_h_tall: float = SIDE_H_TALL
    side_ft_h: float = SIDE_FT_H

    panel_h: float = PANEL_H
    panel_gap: float = PANEL_GAP


def compute_xmax(layout: Layout) -> float:
    return (
        layout.margin_x
        + layout.flow_w
        + layout.gap_flow_side
        + layout.side_w
        + layout.gap_side_panel
        + layout.panel_w
        + layout.margin_x
    )


def required_ymax(layout: Layout) -> float:
    """Min Y span needed to fit flow stack + panel stack."""
    flow_dy = layout.flow_h + layout.flow_v_gap
    flow_height = layout.bottom_margin + layout.top_space + layout.flow_h_tall + 7 * flow_dy

    panel_height = layout.bottom_margin + layout.top_space + (3 * layout.panel_h + 2 * layout.panel_gap)
    return max(flow_height, panel_height, 10.0)


# --------------------------- Drawing ---------------------------

def draw_side_panel(
    ax,
    box,
    title: str,
    body: str,
    c0: str,
    c1: str,
    *,
    body_ha: str = "left",
    body_va: str = "top",
):
    border = "#1f1f1f"
    x, y, w, h = box
    patch = gradient_box(ax, x, y, w, h, c0, c1, edge=border, lw=1.0, r=0.16, shadow=True)

    # title (TOP-CENTER, clipped to panel)
    t_title = ax.text(
        x + w / 2,
        y + h - 0.30,
        title,
        ha="center",
        va="top",
        fontsize=11.2,
        fontweight="bold",
        color="#111111",
        zorder=10,
    )
    t_title.set_clip_path(patch)
    t_title.set_clip_on(True)

    # body (auto-fit; also clipped to panel)
    body_box = (x + 0.02, y + 0.02, w - 0.04, h - 0.55)
    label_in_box(
        ax,
        body_box,
        body,
        ha=body_ha,
        va=body_va,
        max_size=PANEL_BODY_MAX,
        min_size=PANEL_BODY_MIN,
        weight="regular",
        color="#222222",
        pad=0.32,
        clip_patch=patch,
    )

def draw_prisma(ax, c: PrismaCounts, layout: Layout, subtitle: Optional[str] = None, bd_summary: Optional[Dict[str, Any]] = None):
    ax.set_axis_off()

    if AUTO_XMAX:
        layout.xmax = compute_xmax(layout)
    else:
        if layout.xmax <= 0:
            layout.xmax = compute_xmax(layout)

    y_max = required_ymax(layout)
    ax.set_xlim(0, layout.xmax)
    ax.set_ylim(0, y_max)

    border = "#1f1f1f"
    arrowc = "#262626"

    g_id = ("#ffffff", "#dfeeff")
    g_sc = ("#ffffff", "#fff0cc")
    g_el = ("#ffffff", "#e2f7ee")
    g_in = ("#ffffff", "#ffe4ef")
    g_sd = ("#ffffff", "#f0f0f0")

    # title
    title_y = y_max - 0.28
    subtitle_y = y_max - 0.62
    label(ax, layout.xmax / 2, title_y, "PRISMA flow diagram", size=TITLE_SIZE, weight=TITLE_WEIGHT)
    if subtitle:
        label(ax, layout.xmax / 2, subtitle_y, subtitle, size=10.8, color="#444444")

    # column x positions
    xM = layout.margin_x
    xS = xM + layout.flow_w + layout.gap_flow_side
    xP = xS + layout.side_w + layout.gap_side_panel

    wM, wS, wP = layout.flow_w, layout.side_w, layout.panel_w

    # left spine y positions
    h = layout.flow_h
    h_tall = layout.flow_h_tall
    dy = h + layout.flow_v_gap
    y0 = y_max - layout.top_space - h_tall

    b_ident  = (xM, y0,                 wM, h_tall)
    b_unique = (xM, y0 - 1 * dy,        wM, h)
    b_screen = (xM, y0 - 2 * dy,        wM, h)
    b_sought = (xM, y0 - 3 * dy,        wM, h)
    b_retr   = (xM, y0 - 4 * dy,        wM, h)
    b_assess = (xM, y0 - 5 * dy,        wM, h)
    b_incl   = (xM, y0 - 6 * dy,        wM, h)
    b_meta   = (xM, y0 - 7 * dy,        wM, h)

    # right PRISMA boxes (aligned to left)
    b_removed = (xS, y0 - 0.42 * dy,     wS, layout.side_h_tall)
    b_excl_ta = (xS, b_screen[1],        wS, layout.side_h)
    b_notret  = (xS, b_sought[1],        wS, layout.side_h)
    b_excl_ft = (xS, b_assess[1] - 0.62, wS, layout.side_ft_h)

    # rightmost panels (stacked top->bottom)
    panel_top = b_ident[1] + b_ident[3]  # align to top of identification box
    p1 = (xP, panel_top - layout.panel_h,                 wP, layout.panel_h)
    p2 = (xP, panel_top - (2 * layout.panel_h + layout.panel_gap), wP, layout.panel_h)
    p3 = (xP, panel_top - (3 * layout.panel_h + 2 * layout.panel_gap), wP, layout.panel_h)

    # draw left spine boxes
    gradient_box(ax, *b_ident,  *g_id, edge=border, lw=1.32)
    gradient_box(ax, *b_unique, *g_id, edge=border, lw=1.22)
    gradient_box(ax, *b_screen, *g_sc, edge=border, lw=1.32)
    gradient_box(ax, *b_sought, *g_sc, edge=border, lw=1.22)
    gradient_box(ax, *b_retr,   *g_el, edge=border, lw=1.22)
    gradient_box(ax, *b_assess, *g_el, edge=border, lw=1.32)
    gradient_box(ax, *b_incl,   *g_in, edge=border, lw=1.32)
    gradient_box(ax, *b_meta,   *g_in, edge=border, lw=1.22)

    # draw right PRISMA boxes
    gradient_box(ax, *b_removed, *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_excl_ta, *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_notret,  *g_sd, edge=border, lw=1.05)
    gradient_box(ax, *b_excl_ft, *g_sd, edge=border, lw=1.05)

    # section labels (far left)
    label(ax, xM - 0.55, b_ident[1] + b_ident[3] / 2, "Identification", ha="right", size=10.8, weight="bold", color="#2b4d8a")
    label(ax, xM - 0.55, b_screen[1] + h / 2,         "Screening",      ha="right", size=10.8, weight="bold", color="#8a5a00")
    label(ax, xM - 0.55, b_assess[1] + h / 2,         "Eligibility",    ha="right", size=10.8, weight="bold", color="#1f6b3b")
    label(ax, xM - 0.55, b_meta[1] + h / 2,           "Included",       ha="right", size=10.8, weight="bold", color="#7a1f4c")

    # identification box text (slightly more professional)
    ident_lines = [
        "Records identified",
        "• " + _line("Database: PubMed", c.db, zero_kind="searched"),
        "• " + _line("Trial registers", c.reg, zero_kind="searched"),
        "• " + _line("Websites", c.web, zero_kind="searched"),
        "• " + _line("Organisations", c.org, zero_kind="searched"),
        "• " + _line("Citation searching", c.cite, zero_kind="performed"),
    ]
    label_in_box(ax, b_ident, "\n".join(ident_lines), size=BOX_TEXT_SIZE, weight="bold")

    # unique
    unique_n = max(0, (c.db + c.reg + c.web + c.org + c.cite) - (c.dup + c.auto + c.other_rm))
    label_in_box(ax, b_unique, f"Unique records after removals\n(n = {unique_n})",
                 max_size=FLOW_MAX, min_size=FLOW_MIN, weight="bold")

    # screening
    label_in_box(ax, b_screen, f"Records screened (title/abstract)\n(n = {c.screened})",
                 max_size=FLOW_MAX, min_size=FLOW_MIN, weight="bold")

    # sought/retrieved/assessed
    label_in_box(ax, b_sought, f"Reports sought for retrieval (full text)\n(n = {c.reports_sought})",
                 max_size=FLOW_MAX, min_size=FLOW_MIN, weight="bold")

    retrieved_n = max(0, c.reports_sought - c.not_retrieved)
    label_in_box(ax, b_retr, f"Reports retrieved (full text)\n(n = {retrieved_n})",
                 max_size=FLOW_MAX, min_size=FLOW_MIN, weight="bold")

    label_in_box(ax, b_assess, f"Reports assessed for eligibility (full text)\n(n = {c.assessed})",
                 max_size=FLOW_MAX, min_size=FLOW_MIN, weight="bold")

    # included
    incl_txt = (
        f"Studies included in review (qualitative synthesis) (n = {c.studies})\n"
        f"Reports included (n = {c.reports})"
    )
    label_in_box(ax, b_incl, incl_txt, wrap=False, size=BOX_TEXT_SIZE, weight="bold")


    # meta
    if c.meta_studies > 0:
        meta_txt = f"Studies included in meta-analysis (quantitative)\n(n = {c.meta_studies})"
    else:
        meta_txt = "Meta-analysis not performed\n(n = 0)"
    label_in_box(ax, b_meta, meta_txt, max_size=10.6, min_size=8.0, weight="bold")

    # right PRISMA text
    auto_state = "not used" if c.auto == 0 else "used"
    other_line = (
        f"Other reasons not performed (n = {c.other_rm})"
        if c.other_rm == 0
        else f"Other reasons (n = {c.other_rm})"
    )

    removed_txt = "\n".join(
        [
            "Records removed before screening",
            f"Duplicate records removed (n = {c.dup})",
            f"Automation tools {auto_state} (n = {c.auto})",
            other_line,
        ]
    )
    label_in_box(
        ax, b_removed, removed_txt,
        ha="center", va="center",
        justify=False, wrap=True, size=BOX_TEXT_SIZE
    )

    label_in_box(ax, b_excl_ta, f"Records excluded\n(n = {c.excluded_ta})",
                 max_size=SIDE_MAX, min_size=SIDE_MIN, weight="bold")

    label_in_box(ax, b_notret, f"Reports not retrieved\n(n = {c.not_retrieved})",
                 max_size=SIDE_MAX, min_size=SIDE_MIN, weight="bold")

    if c.fulltext_excl:
        lines = ["Reports excluded (with reasons)"]
        for r, n in c.fulltext_excl:
            r = re.sub(r"\s+", " ", str(r)).strip()

            # Shorten the long reason to keep the line compact in the box
            if re.search(r"wrong outcome", r, flags=re.IGNORECASE) and re.search(r"non[- ]?MM", r, flags=re.IGNORECASE):
                r = "Non-MM plasma cell disorder"

            lines.append(f"{r} (n = {n})")
        ft_txt = "\n".join(lines)

    label_in_box(
        ax, b_excl_ft, ft_txt,
        ha="center", va="center",
        justify=False, wrap=True,
        size=BOX_TEXT_SIZE
    )

    # arrows: left spine
    connect_down(ax, b_ident, b_unique, color=arrowc)
    connect_down(ax, b_unique, b_screen, color=arrowc)
    connect_down(ax, b_screen, b_sought, color=arrowc)
    connect_down(ax, b_sought, b_retr, color=arrowc)
    connect_down(ax, b_retr, b_assess, color=arrowc)
    connect_down(ax, b_assess, b_incl, color=arrowc)
    connect_down(ax, b_incl, b_meta, color=arrowc)

    # elbow arrows to right PRISMA boxes
    elbow = "angle3,angleA=0,angleB=90"
    arrow(ax, (xM + wM + 0.12, b_ident[1] + b_ident[3] / 2), (xS - 0.10, b_removed[1] + b_removed[3] / 2),
          color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_screen[1] + h / 2), (xS - 0.10, b_excl_ta[1] + b_excl_ta[3] / 2),
          color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_sought[1] + h / 2), (xS - 0.10, b_notret[1] + b_notret[3] / 2),
          color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)
    arrow(ax, (xM + wM + 0.12, b_assess[1] + h / 2), (xS - 0.10, b_excl_ft[1] + b_excl_ft[3] / 2),
          color=arrowc, lw=1.25, ms=12.0, connectionstyle=elbow)

        # rightmost panels content
    # Search snapshot (keep it readable — no raw PubMed syntax dump)
    search_lines: List[str] = []
    if c.search_date:
        search_lines.append(f"Last search date: {c.search_date}")
    if c.databases_list:
        search_lines.append("Database(s): " + ", ".join(c.databases_list))
    if c.limits:
        search_lines.append("Limits: " + c.limits)

    strategy = (
        "Search strategy:\n"
        "Condition: multiple myeloma / plasma cell myeloma / plasma cell neoplasm*\n"
        "Location: Bangladesh / Bangladeshi\n"
        "Fields: MeSH terms; title/abstract"
    )
    search_body = "\n".join(search_lines + [strategy]) if search_lines else strategy
    draw_side_panel(ax, p1, "Search snapshot", search_body, "#ffffff", "#eef5ff", body_ha="center", body_va="center")

    # Bangladesh evidence summary (from DOCX if available)
    if bd_summary:
        studies = bd_summary.get("studies")
        total_n = bd_summary.get("total_n")
        n_counted = bd_summary.get("n_counted")
        designs = bd_summary.get("designs", {}) or {}

        lines: List[str] = []
        if studies is not None:
            lines.append(f"Extracted Table 1 studies: {studies}")
        if total_n is not None and n_counted:
            lines.append(f"Sum sample size (extractable): {total_n} (from {n_counted} studies)")
        if designs:
            lines.append("Study designs:")
            for k, v in designs.items():
                k = re.sub(r"\s+", " ", str(k)).strip()
                lines.append(f"{k}: {v}")

        bd_body = "\n".join(lines) if lines else "Evidence summary available, but no extractable fields found"
    else:
        bd_body = "Evidence summary not available\n(Table 1 not found in DOCX)"

    draw_side_panel(ax, p2, "Bangladesh evidence summary", bd_body, "#ffffff", "#eefcf2", body_ha="center", body_va="center")

    # Diagnostics
    diag = (
        "Consistency checks (recommended):\n"
        "• verify arithmetic totals\n"
        "• confirm exclusions match reasons\n"
        "• confirm retrieval counts"
    )
    draw_side_panel(ax, p3, "Diagnostics", diag, "#ffffff", "#fff0f4", body_ha="center", body_va="center")

    return {"panel_boxes": (p1, p2, p3)}


def write_template(path: Path):
    tpl = {
        "records_identified": {"databases": 0, "registers": 0},
        "records_identified_other_methods": {"websites": 0, "organisations": 0, "citation_searching": 0},
        "records_removed_before_screening": {"duplicates": 0, "automation": 0, "other": 0},
        "records_screened": 0,
        "records_excluded": 0,
        "reports_sought_for_retrieval": 0,
        "reports_not_retrieved": 0,
        "reports_assessed_for_eligibility": 0,
        "reports_excluded": [
            {"reason": "Not Bangladesh population", "n": 0},
            {"reason": "Non-MM plasma cell disorder", "n": 0},
            {"reason": "Case report or non-eligible design", "n": 0},
            {"reason": "Insufficient extractable data", "n": 0},
        ],
        "studies_included_in_review": 0,
        "reports_included_in_review": 0,
        "studies_included_in_meta_analysis": 0,
        "metadata": {
            "last_search_date": "DD-MMM-YYYY",
            "databases": ["PubMed"],
            "limits": "No filters/limits applied",
            "notes": "Insert the exact PubMed query string here.",
        },
    }
    path.write_text(json.dumps(tpl, indent=2), encoding="utf-8")


def main():
    set_rc()
    ap = argparse.ArgumentParser()
    ap.add_argument("--counts", type=str, default=None)
    ap.add_argument("--write_template", type=str, default=None)
    ap.add_argument("--docx", type=str, default=None)
    ap.add_argument("--outdir", type=str, default="figures")
    ap.add_argument("--basename", type=str, default="Figure_1_PRISMA")
    ap.add_argument("--dpi", type=int, default=600)
    ap.add_argument("--eps", action="store_true")
    ap.add_argument("--subtitle", type=str, default=None)
    args = ap.parse_args()

    if args.write_template:
        p = Path(args.write_template)
        p.parent.mkdir(parents=True, exist_ok=True)
        write_template(p)
        print(f"Wrote template to: {p}")
        return

    if not args.counts:
        raise SystemExit("ERROR: provide --counts prisma_counts.json (or create one with --write_template).")

    counts_path = Path(args.counts)
    if not counts_path.exists():
        raise SystemExit(f"ERROR: counts file not found: {counts_path}")

    d = json.loads(counts_path.read_text(encoding="utf-8"))
    c = PrismaCounts.from_json(d)

    bd_summary = None
    if args.docx:
        bd_summary = extract_table1_summary(Path(args.docx))

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    layout = Layout()

    # Figure size: keep the same visual scale; layout.ymax drives axes limits
    fig = plt.figure(figsize=(18.6, 8.8))
    ax = fig.add_subplot(1, 1, 1)
    info = draw_prisma(ax, c, layout=layout, subtitle=args.subtitle, bd_summary=bd_summary)


    base = outdir / args.basename
    save_kws = dict(bbox_inches="tight", facecolor="white", pad_inches=0.20)
    fig.savefig(str(base) + ".png", dpi=args.dpi, **save_kws)
    fig.savefig(str(base) + ".pdf", **save_kws)
    fig.savefig(str(base) + ".svg", **save_kws)
    if args.eps:
        fig.savefig(str(base) + ".eps", **save_kws)
    plt.close(fig)

    print(f"Saved: {base}.png/.pdf/.svg" + ("/.eps" if args.eps else ""))


if __name__ == "__main__":
    main()
