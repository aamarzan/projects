# make_table4_premium_docx.py
# Output: Table_4_BD_Minimum_Dataset_and_Gaps_Premium.docx

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from pathlib import Path

def set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for key in ("sz", "val", "color", "space"):
                if key in edge_data:
                    tag.set(qn(f"w:{key}"), str(edge_data[key]))
            tcBorders.append(tag)

def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

headers = [
    "Domain",
    "Evidence backed critical points from the reference set",
    "Standardized variables and definitions to extract",
    "Bangladesh reporting status across BD studies",
    "Critical gaps to address (for Q1 readiness)",
    "Bangladesh feasible minimum dataset for future cohorts",
    "Key supporting refs (your list #)",
]

rows = [
    ("Case definition and disease spectrum",
     "Spectrum from precursor conditions to overt multiple myeloma; diagnostic nuance in borderline marrow plasma cell percentages.",
     "Criteria source; disease state (MM vs MGUS vs SMM); marrow plasma cell percent; defining events; diagnosis date; new vs relapsed.",
     "Often no explicit criteria reference; precursor states not separated.",
     "Misclassification risk and low comparability.",
     "Criteria reference; marrow percent; defining events; diagnosis date; new vs relapsed field.",
     "1,3,58"),
    ("Diagnostic pathway and baseline workup",
     "Baseline workup should be consistent; atypical presentation risk exists.",
     "CBC; renal tests; calcium; SPEP; immunofixation; urine protein method; marrow method; imaging type; per patient completeness.",
     "Tests reported but assay methods and completeness vary.",
     "Not reproducible pathway synthesis.",
     "Per patient checklist + completeness score; record assay method if stated.",
     "1,2,38,40,41,46"),
    ("Staging and risk stratification",
     "Staging uses biomarkers; revised staging needs LDH and cytogenetics.",
     "ISS inputs; LDH; cytogenetics availability; R-ISS feasibility documentation.",
     "ISS sometimes reported; LDH and cytogenetics often missing.",
     "Modern risk reporting not possible.",
     "Albumin; beta2 microglobulin; LDH; cytogenetics attempted yes/no plus reason.",
     "47,49,50,52,53,54,23"),
    ("Monoclonal protein characterization and immunotype",
     "Protein type relates to phenotype and renal risk including light chain disease and proteinuria.",
     "SPEP yes/no; immunofixation yes/no; Ig type; light chain type; FLC availability; urine test method.",
     "Monoclonal band common; immunofixation and FLC inconsistently described.",
     "Weak renal and relapse interpretation.",
     "SPEP + immunofixation + urine method; FLC availability field.",
     "21,56,22,46"),
    ("Bone disease and skeletal complications",
     "Bone resorption mechanism; pain and bone management are essential.",
     "Imaging modality; lytic criteria; fracture definition; vertebral collapse definition; pain measurement.",
     "Lesions reported; imaging criteria vary.",
     "Heterogeneity blocks pooling.",
     "Baseline imaging type; sites assessed; standard skeletal related event definition.",
     "20,28,34,35,30"),
    ("Renal involvement and classification",
     "Renal disease not captured by creatinine alone; classification matters.",
     "Renal definition; proteinuria; lesion type if evaluated; dialysis or supportive measures.",
     "Mostly creatinine threshold only.",
     "Missing renal phenotyping.",
     "Creatinine; eGFR if possible; urine protein method; cast nephropathy evaluation yes/no.",
     "22,43,44,40,41,46"),
    ("Anemia evaluation and management",
     "Anemia clinically important; evaluation and management matter.",
     "Anemia definition; workup elements; transfusion; ESA use.",
     "Hb reported; workup and management usually missing.",
     "Cannot propose practical pathway.",
     "Hb; transfusion events; ESA use; workup performed yes/no.",
     "24,36,37"),
    ("Infection risk and prevention",
     "Immunodeficiency and prevention key; IVIG sometimes considered.",
     "Infection definitions; prophylaxis; vaccination; IVIG use and availability.",
     "Prevention rarely standardized.",
     "High impact actionable gap.",
     "Severe infection events; prophylaxis and vaccine advice; IVIG availability field.",
     "32,55"),
    ("Early detection and referral triggers",
     "Earlier detection may improve survival; primary care blood tests can flag MM.",
     "Symptom onset; first contact; trigger labs; referral pathway; missed opportunities.",
     "Delay metrics rarely reported.",
     "Cannot build BD referral pathway recommendations.",
     "Diagnostic delay dataset plus trigger test list.",
     "3,16,38,39,42"),
    ("Biomarkers beyond routine",
     "Emerging biomarkers are future direction; classic prognostic markers support staging logic.",
     "Non-routine biomarkers tested; intended role; feasibility barriers.",
     "Mostly absent.",
     "Should be framed as future research priority.",
     "Minimum: LDH plus cytogenetics availability before advanced biomarkers.",
     "45,47,49,23"),
    ("Sex differences",
     "Outcomes may differ by sex; report stratified.",
     "Sex stratified outcomes; sex in adjusted analysis.",
     "Mostly only male predominance reported.",
     "Weak modern reporting.",
     "Outcome stratified by sex; include sex as covariate.",
     "15,31"),
    ("Elderly and vulnerability",
     "Very elderly need different approach; subgroup reporting matters.",
     "Age strata; ECOG; comorbidity; treatment modification by age.",
     "Age and sometimes ECOG reported; frailty indices rare.",
     "Limited clinical applicability.",
     "ECOG; comorbidity count; age strata.",
     "33,51,29"),
    ("Treatment documentation and access",
     "Era and access drive outcomes; BD context is important.",
     "Regimen class; cycles; ASCT availability; maintenance; access barriers.",
     "Regimen reporting inconsistent.",
     "Cannot interpret outcome differences.",
     "Regimen; cycles; drug availability; ASCT availability; barrier yes/no.",
     "5,6,9,10,13"),
    ("Outcome reporting and follow up",
     "Endpoints and follow up must be explicit for synthesis.",
     "Follow up duration; endpoints defined; response criteria; early mortality definition.",
     "Endpoints inconsistently defined.",
     "Blocks meta-analysis.",
     "Follow up duration; vital status fixed timepoints; early mortality definition.",
     "7,9,59"),
]

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = Inches(0.6)
sec.right_margin = Inches(0.6)
sec.top_margin = Inches(0.6)
sec.bottom_margin = Inches(0.6)

title = doc.add_paragraph("Table 4. Bangladesh relevant minimum dataset and evidence gap matrix derived from the reference set.")
title.runs[0].font.name = "Calibri"
title.runs[0].font.size = Pt(14)
title.runs[0].font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Purpose: define standardized variables needed for comparability and meta analysis readiness, without duplicating Tables 1 and 2.")
subtitle.runs[0].font.name = "Calibri"
subtitle.runs[0].font.size = Pt(10)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

table = doc.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

col_widths = [Cm(3.2), Cm(8.4), Cm(8.0), Cm(6.0), Cm(6.6), Cm(8.0), Cm(3.2)]

hdr = table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].width = col_widths[i]
    hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(hdr[i], "1F4E79")
    set_cell_borders(hdr[i],
        top={"sz": 10, "val": "single", "color": "1F4E79"},
        bottom={"sz": 10, "val": "single", "color": "1F4E79"},
        left={"sz": 10, "val": "single", "color": "1F4E79"},
        right={"sz": 10, "val": "single", "color": "1F4E79"},
    )
set_repeat_table_header(table.rows[0])

for idx, row in enumerate(rows, start=1):
    cells = table.add_row().cells
    shade = "FFFFFF" if idx % 2 else "F2F2F2"
    for cidx, val in enumerate(row):
        cells[cidx].width = col_widths[cidx]
        cells[cidx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cells[cidx].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        rr = p.add_run(val)
        rr.font.name = "Calibri"
        rr.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cells[cidx], shade)
        set_cell_borders(cells[cidx],
            top={"sz": 6, "val": "single", "color": "D9D9D9"},
            bottom={"sz": 6, "val": "single", "color": "D9D9D9"},
            left={"sz": 6, "val": "single", "color": "D9D9D9"},
            right={"sz": 6, "val": "single", "color": "D9D9D9"},
        )

doc.add_paragraph("")
note = doc.add_paragraph("Note: This table focuses on standardization and evidence gaps. Descriptive prevalence values are summarized in Tables 1 and 2.")
note.runs[0].font.name = "Calibri"
note.runs[0].font.size = Pt(9)

abbr = doc.add_paragraph("Abbreviations: ASCT autologous stem cell transplantation; CBC complete blood count; ECOG Eastern Cooperative Oncology Group; ISS International Staging System; R-ISS Revised International Staging System; ESA erythropoiesis stimulating agent.")
abbr.runs[0].font.name = "Calibri"
abbr.runs[0].font.size = Pt(9)

out = Path("Table_4_BD_Minimum_Dataset_and_Gaps_Premium.docx")
doc.save(out)
print("Saved:", out.resolve())
