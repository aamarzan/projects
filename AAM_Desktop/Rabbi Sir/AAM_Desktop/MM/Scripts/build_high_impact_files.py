#!/usr/bin/env python3
"""
build_high_impact_files.py

Creates 4 high-impact submission files (DOCX) + CSV templates for figures:

1) PRISMA_2020_Checklist_Bangladesh_MM.docx
2) Appendix_Search_Strategy_Bangladesh_MM.docx
3) Supplement_Risk_of_Bias_Bangladesh_MM.docx
4) Data_Code_Availability_and_File_Manifest_Bangladesh_MM.docx

Also creates:
- rob_template.csv
- forest_template.csv
- evidence_matrix_auto.csv  (auto-filled from Table 01 "Key Findings" keyword scan)
"""

from __future__ import annotations
import argparse, re, csv
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


# -----------------------------
# Helpers
# -----------------------------
def _set_margins(doc: Document, top=0.8, bottom=0.8, left=0.7, right=0.7):
    sec = doc.sections[0]
    sec.top_margin = Inches(top)
    sec.bottom_margin = Inches(bottom)
    sec.left_margin = Inches(left)
    sec.right_margin = Inches(right)

def _title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)

def _note(doc: Document, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)

def _cell(cell, text: str, bold=False, size=9, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def read_table01_studies(table01_docx: Path):
    """
    Expects Table 01 with headers like:
    Study Reference | Setting | Study Period | Method | Sample size | Key Findings
    """
    d = Document(str(table01_docx))
    if not d.tables:
        raise ValueError("No table found in Table 01 DOCX.")
    t = d.tables[0]
    rows = [[c.text.strip() for c in r.cells] for r in t.rows]
    header = rows[0]

    # find best-effort indices
    def find_col(name_candidates):
        for cand in name_candidates:
            for i, h in enumerate(header):
                if h.strip().lower() == cand.lower():
                    return i
        # fallback partial match
        for cand in name_candidates:
            for i, h in enumerate(header):
                if cand.lower() in h.strip().lower():
                    return i
        return None

    idx_study = find_col(["Study Reference", "Study", "Reference"])
    idx_method = find_col(["Method", "Study design", "Design"])
    idx_key    = find_col(["Key Findings", "Key finding", "Findings"])
    idx_n      = find_col(["Sample size", "N", "Sample"])

    if idx_study is None:
        idx_study = 0

    studies = []
    for r in rows[1:]:
        if len(r) <= idx_study:
            continue
        s = r[idx_study].strip()
        if not s:
            continue
        studies.append({
            "study": s,
            "method": (r[idx_method].strip() if idx_method is not None and idx_method < len(r) else ""),
            "n": (r[idx_n].strip() if idx_n is not None and idx_n < len(r) else ""),
            "key": (r[idx_key].strip() if idx_key is not None and idx_key < len(r) else ""),
        })
    return studies


# -----------------------------
# 1) PRISMA checklist DOCX
# -----------------------------
def make_prisma_checklist(outdir: Path):
    doc = Document()
    _set_margins(doc)
    _title(doc, "PRISMA 2020 Checklist (Completed) — Bangladesh Multiple Myeloma Systematic Review")
    _note(doc, "Tip: Update page numbers after final formatting; locations below use section/heading pointers.")

    items = [
        ("TITLE", "1", "Title", "Identify the report as a systematic review."),
        ("ABSTRACT", "2", "Abstract", "See the PRISMA 2020 for Abstracts checklist."),
        ("INTRODUCTION", "3", "Rationale", "Describe the rationale for the review in the context of existing knowledge."),
        ("INTRODUCTION", "4", "Objectives", "Provide explicit objective(s) or question(s) the review addresses."),
        ("METHODS", "5", "Eligibility criteria", "Specify inclusion/exclusion criteria; how studies were grouped for syntheses."),
        ("METHODS", "6", "Information sources", "Specify all sources searched and last search date for each."),
        ("METHODS", "7", "Search strategy", "Present full search strategies, including filters/limits."),
        ("METHODS", "8", "Selection process", "Specify screening methods, reviewers, independence, and automation."),
        ("METHODS", "9", "Data collection process", "Specify how data were extracted, reviewers, independence, verification."),
        ("METHODS", "10a", "Data items", "List and define outcomes for which data were sought."),
        ("METHODS", "10b", "Data items", "List and define other variables; assumptions for missing/unclear info."),
        ("METHODS", "11", "Risk of bias assessment", "Methods to assess risk of bias in included studies (tool, reviewers, etc.)."),
        ("METHODS", "12", "Effect measures", "Effect measures used for each outcome (e.g., prevalence, mean/median)."),
        ("METHODS", "13", "Synthesis methods", "Describe synthesis approach; if meta-analysis, model/heterogeneity/software."),
        ("METHODS", "14", "Reporting bias assessment", "Methods to assess missing results/reporting bias (if applicable)."),
        ("METHODS", "15", "Certainty assessment", "Methods to assess certainty/confidence (e.g., GRADE) (if applicable)."),
        ("RESULTS", "16", "Study selection", "Numbers screened/assessed/included; flow diagram."),
        ("RESULTS", "17", "Study characteristics", "Cite each included study and present characteristics."),
        ("RESULTS", "18", "Risk of bias in studies", "Present per-study RoB judgments."),
        ("RESULTS", "19", "Results of individual studies", "Present outcome results per study (as available)."),
        ("RESULTS", "20", "Results of syntheses", "If synthesis/meta-analysis performed, present pooled results/heterogeneity."),
        ("DISCUSSION", "23", "Discussion", "Interpretation, limitations (evidence + process), implications."),
        ("OTHER", "24", "Registration/protocol", "Registration, protocol access, and amendments (or state not registered)."),
        ("OTHER", "25", "Support", "Funding/support and role of funders."),
        ("OTHER", "26", "Competing interests", "COI declaration."),
        ("OTHER", "27", "Availability", "Data/code/material availability and where."),
    ]

    # Best-effort “location pointers” for your current manuscript
    loc = {
        "1": "Title",
        "2": "ABSTRACT",
        "3": "INTRODUCTION",
        "4": "ABSTRACT (Objective) + INTRODUCTION",
        "5": "METHODOLOGY (Inclusion/Exclusion criteria)",
        "6": "Methods + Figure 1 (Search snapshot)",
        "7": "Figure 1 + Appendix (this file set)",
        "8": "Methods (add: reviewer workflow)",
        "9": "Methods (add: extraction workflow)",
        "10a": "Methods (add explicit outcome list)",
        "10b": "Methods (data items sentence)",
        "11": "ADD (Supplement: RoB table + plot)",
        "12": "Methods/Results (state measures: prevalence/range)",
        "13": "Methods/Results (narrative synthesis; no meta-analysis)",
        "14": "N/A if no meta-analysis (state explicitly)",
        "15": "Optional (GRADE) or justify not assessed",
        "16": "Figure 1 PRISMA flow",
        "17": "Table 01",
        "18": "ADD (RoB supplement)",
        "19": "Results + Table 02 (optional per-study table)",
        "20": "N/A if no meta-analysis (or optional forest without pooling)",
        "23": "DISCUSSION + LIMITATION + CONCLUSION",
        "24": "ADD (Methods end)",
        "25": "ADD (Funding/Acknowledgements)",
        "26": "ADD (Competing interests)",
        "27": "ADD (Data/Code availability statement)",
    }

    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Section", "Item #", "Checklist item", "Location", "Action needed"]):
        _cell(hdr[i], h, bold=True, size=10, align="center")

    # Minimal “action” suggestions (journal IF ~5)
    action_map = {
        "2": "Ensure PRISMA-Abstracts elements (databases, dates, eligibility, synthesis type).",
        "6": "State all sources (PubMed + handsearch) and last search date per source.",
        "7": "Provide exact query strings in Appendix.",
        "8": "Specify #screeners, independent screening, tie-breaker.",
        "9": "Specify extraction form + double-checking.",
        "10a": "List outcomes explicitly (age/sex, symptoms, labs, imaging, ISS, etc.).",
        "11": "Add RoB tool + per-study judgments + summary plot.",
        "24": "State protocol registration (or explicitly 'not registered').",
        "25": "Add funding/support statement (even if none).",
        "26": "Add COI statement.",
        "27": "Add data/code/material availability statement.",
    }

    for section, itemno, topic, text in items:
        r = tbl.add_row().cells
        _cell(r[0], section)
        _cell(r[1], itemno, align="center")
        _cell(r[2], f"{topic}: {text}")
        _cell(r[3], loc.get(itemno, "—"))
        _cell(r[4], action_map.get(itemno, "—"))

    doc.add_paragraph()
    _note(doc, "Reference: Page MJ et al. PRISMA 2020. BMJ 2021;372:n71.")
    out = outdir / "PRISMA_2020_Checklist_Bangladesh_MM.docx"
    doc.save(str(out))


# -----------------------------
# 2) Search strategy appendix DOCX
# -----------------------------
def make_search_appendix(outdir: Path, last_search_date: str, query: str):
    doc = Document()
    _set_margins(doc)
    _title(doc, "Supplementary Appendix — Reproducible Search Strategy")
    _note(doc, "Use this appendix to satisfy PRISMA Items 6–7 (Information sources and Search strategy).")

    doc.add_heading("1. Search snapshot", level=1)
    doc.add_paragraph(f"Last search date: {last_search_date}")
    doc.add_paragraph("Database(s): PubMed (National Library of Medicine)")
    doc.add_paragraph("Limits: No filters/limits applied (all years; all languages; all article types).")

    doc.add_heading("2. PubMed query (as run; no filters)", level=1)
    p = doc.add_paragraph()
    run = p.add_run(query)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("3. Additional identification methods", level=1)
    doc.add_paragraph("• Citation searching / reference list checking of eligible and closely related articles.", style=None)
    doc.add_paragraph("• Trial registers/websites/organisations: not searched (state if unchanged).", style=None)

    doc.add_heading("4. De-duplication and screening workflow (to paste into Methods)", level=1)
    doc.add_paragraph(
        "Duplicates were removed before screening. Title/abstract screening was performed on unique records, "
        "followed by full-text assessment. Automation tools were not used. "
        "Add: number of reviewers, independence, and tie-break process."
    )

    doc.add_paragraph()
    _note(doc, "IMPORTANT consistency check: If Methods says 'till Nov 2025' but the search date is later, align them.")
    out = outdir / "Appendix_Search_Strategy_Bangladesh_MM.docx"
    doc.save(str(out))


# -----------------------------
# 3) Risk of bias supplement DOCX + rob_template.csv
# -----------------------------
ROB_DOMAINS = [
    "D1 Sample frame appropriate",
    "D2 Sampling method",
    "D3 Sample size adequate",
    "D4 Subjects/setting described",
    "D5 Coverage of sample in analysis",
    "D6 Valid condition identification",
    "D7 Standard measurement for all",
    "D8 Appropriate statistical analysis",
    "D9 Response rate adequate/managed",
]

def make_rob_supplement(outdir: Path, studies):
    doc = Document()
    _set_margins(doc)
    _title(doc, "Supplementary — Risk of Bias / Quality Appraisal (Study-level)")
    _note(doc, "Recommended tool: JBI Critical Appraisal Checklist for Studies Reporting Prevalence Data (domains below).")
    _note(doc, "Fill judgments using full-texts. If full-text is not available, mark 'Unclear' and state why.")

    doc.add_heading("A) Per-study risk of bias table (fill/verify)", level=1)
    tbl = doc.add_table(rows=1, cols=3 + len(ROB_DOMAINS))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    _cell(hdr[0], "Study", bold=True, align="center")
    _cell(hdr[1], "Design", bold=True, align="center")
    _cell(hdr[2], "Overall", bold=True, align="center")
    for i, d in enumerate(ROB_DOMAINS, start=3):
        _cell(hdr[i], d, bold=True, align="center", size=9)

    for s in studies:
        r = tbl.add_row().cells
        _cell(r[0], s["study"])
        _cell(r[1], s["method"] or "—", align="center")
        _cell(r[2], "Unclear", align="center")
        for j in range(len(ROB_DOMAINS)):
            _cell(r[3 + j], "Unclear", align="center")

    doc.add_paragraph()
    doc.add_heading("B) Figure generation", level=1)
    doc.add_paragraph("Use make_rob_summary_plot.py with rob_template.csv to generate:")
    doc.add_paragraph("• Traffic-light plot (study × domain)")
    doc.add_paragraph("• Summary bars (proportion Low/High/Unclear per domain)")

    out_docx = outdir / "Supplement_Risk_of_Bias_Bangladesh_MM.docx"
    doc.save(str(out_docx))

    # CSV template for plotting
    out_csv = outdir / "rob_template.csv"
    with out_csv.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Study", "Design", "Overall"] + ROB_DOMAINS)
        for s in studies:
            w.writerow([s["study"], s["method"] or "", "Unclear"] + ["Unclear"] * len(ROB_DOMAINS))


# -----------------------------
# 4) Data + code availability + manifest DOCX
# -----------------------------
def make_data_code_manifest(outdir: Path, studies):
    doc = Document()
    _set_margins(doc)
    _title(doc, "Data & Code Availability + File Manifest (Submission Package)")

    doc.add_heading("1) Data availability statement (paste-ready)", level=1)
    doc.add_paragraph(
        "The data extracted from included studies and the derived summary tables are provided as Supplementary materials. "
        "If any source data are restricted (e.g., full-text PDFs), this is due to publisher copyright; extracted variables remain available."
    )

    doc.add_heading("2) Code availability statement (paste-ready)", level=1)
    doc.add_paragraph(
        "All analysis and figure-generation scripts used to produce the PRISMA diagram and supplementary figures are provided with the submission "
        "(or are available in a public repository, if you choose to host them)."
    )

    doc.add_heading("3) Materials availability (paste-ready)", level=1)
    doc.add_paragraph(
        "Search export files (date-stamped), screening logs (optional), extraction forms, and quality appraisal forms are available as Supplementary files "
        "or from the corresponding author upon reasonable request."
    )

    doc.add_heading("4) File manifest (what the journal receives)", level=1)
    p = doc.add_paragraph()
    p.add_run("Core figures/tables already prepared:\n").bold = True
    doc.add_paragraph("• Figure 1: PRISMA flow diagram")
    doc.add_paragraph("• Table 01 / Table 02 (main synthesis tables)")
    doc.add_paragraph("• Table 03–04 and Supplementary Tables 1–4 (methods/gaps/governance)")

    p = doc.add_paragraph()
    p.add_run("\nNew high-impact files created by this script:\n").bold = True
    doc.add_paragraph("• PRISMA_2020_Checklist_Bangladesh_MM.docx")
    doc.add_paragraph("• Appendix_Search_Strategy_Bangladesh_MM.docx")
    doc.add_paragraph("• Supplement_Risk_of_Bias_Bangladesh_MM.docx")
    doc.add_paragraph("• Data_Code_Availability_and_File_Manifest_Bangladesh_MM.docx")

    doc.add_heading("5) Included-study list (for cross-check)", level=1)
    for s in studies:
        doc.add_paragraph(f"• {s['study']} ({s['method'] or 'design NR'})")

    out = outdir / "Data_Code_Availability_and_File_Manifest_Bangladesh_MM.docx"
    doc.save(str(out))


# -----------------------------
# Evidence matrix auto CSV
# -----------------------------
HEATMAP_FEATURES = {
    "Age": [r"\bage\b", r"mean age", r"median age"],
    "Sex/Male%": [r"\bmale\b", r"male-to-female", r"\bM:\s*F\b"],
    "Bone pain": [r"bone pain"],
    "Anemia/Hb": [r"\banemi", r"\bhb\b", r"hemoglobin"],
    "Renal impairment/Cr": [r"creatin", r"renal", r"kidney"],
    "Hypercalcemia/Ca": [r"calcium", r"hypercal"],
    "Bence Jones": [r"bence", r"\bBJP\b", r"proteinuria"],
    "Lytic lesions": [r"lytic", r"lesion", r"vertebral", r"skull"],
    "ISS stage": [r"\bISS\b", r"stage III", r"staging"],
    "IgG/Immunofixation": [r"\bIgG\b", r"\bIgA\b", r"immunofix", r"electrophoresis"],
    "β2-microglobulin": [r"microglobulin", r"β2", r"2-microglobulin"],
    "Bone marrow plasma%": [r"bone marrow", r"plasma cell", r"infiltration"],
}

def make_evidence_matrix_auto(outdir: Path, studies):
    out = outdir / "evidence_matrix_auto.csv"
    cols = ["Study"] + list(HEATMAP_FEATURES.keys())
    with out.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(cols)
        for s in studies:
            txt = (s.get("key") or "").lower()
            row = [s["study"]]
            for feat, pats in HEATMAP_FEATURES.items():
                hit = 0
                for pat in pats:
                    if re.search(pat, txt, flags=re.IGNORECASE):
                        hit = 1
                        break
                row.append(hit)
            w.writerow(row)


# -----------------------------
# Main
# -----------------------------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--table01", required=True, help="Path to Table 01 DOCX (study list + key findings).")
    ap.add_argument("--outdir", default="HIGH_IMPACT_FILES", help="Output folder.")
    ap.add_argument("--last_search_date", default="15-Feb-2026", help="Last search date (match PRISMA).")
    ap.add_argument(
        "--pubmed_query",
        default='("Multiple Myeloma"[MeSH Terms] OR "multiple myeloma"[tiab] OR "plasma cell myeloma"[tiab] OR "plasma cell neoplasm*"[tiab]) AND (Bangladesh[MeSH Terms] OR Bangladesh[tiab] OR Bangladeshi[tiab])',
        help="PubMed query as run."
    )
    args = ap.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    studies = read_table01_studies(Path(args.table01))

    make_prisma_checklist(outdir)
    make_search_appendix(outdir, args.last_search_date, args.pubmed_query)
    make_rob_supplement(outdir, studies)
    make_data_code_manifest(outdir, studies)

    # templates for plots
    make_evidence_matrix_auto(outdir, studies)

    # forest template
    forest_csv = outdir / "forest_template.csv"
    with forest_csv.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Outcome", "Study", "Estimate", "CI_low", "CI_high", "N", "Measure"])
        for s in studies:
            w.writerow(["(e.g., Bone pain prevalence)", s["study"], "", "", "", s.get("n",""), "proportion"])

    print(f"\nDone. Created files in: {outdir.resolve()}\n")
    print("Next:")
    print("1) Fill rob_template.csv (Low/High/Unclear/NA).")
    print("2) Fill forest_template.csv with estimates + CIs (optional).")
    print("3) Run plotting scripts (below).")

if __name__ == "__main__":
    main()
